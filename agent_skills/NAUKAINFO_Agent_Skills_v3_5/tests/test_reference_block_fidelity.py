from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / 'scripts'))
import normalize_captions_references as ncr


class ReferenceBlockFidelityTest(unittest.TestCase):
    def test_stamp_spacing_and_fresh_hanging_numbering(self):
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / 'out.docx'
            d = Document()
            existing = {s.name for s in d.styles}
            for name in ['REF-TITLE', 'REFER']:
                if name not in existing:
                    d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            d.add_paragraph('Останній абзац статті.')
            d.add_paragraph('')
            d.add_paragraph('')
            d.add_paragraph('СПИСОК ВИКОРИСТАНОЇ ЛІТЕРАТУРИ')
            d.add_paragraph('')
            d.add_paragraph('Перше джерело з довгим текстом для перенесення рядка.')
            d.add_paragraph('Друге джерело.')
            d.add_paragraph(ncr.TAIL_MARKER)
            audit = {}
            ncr.normalize_reference_block(d, 3, 7, audit)
            d.save(out)
            x = Document(out)
            paras = x.paragraphs
            hi = next(i for i, p in enumerate(paras) if p.text.strip() == ncr.REF_STAMP_UA)
            self.assertEqual(paras[hi].style.name, 'REF-TITLE')
            self.assertFalse(paras[hi - 1].text.strip())
            self.assertTrue(hi < 2 or paras[hi - 2].text.strip())
            self.assertFalse(paras[hi + 1].text.strip())
            refs = [p for p in paras[hi + 2:] if p.style.name == 'REFER']
            self.assertEqual(len(refs), 2)
            nums = set()
            for p in refs:
                pPr = p._p.pPr
                self.assertIsNotNone(pPr.numPr)
                nums.add(pPr.numPr.numId.val)
                self.assertIsNone(pPr.find(qn('w:ind')))
                self.assertIsNone(pPr.find(qn('w:tabs')))
            self.assertEqual(len(nums), 1)
            self.assertGreater(next(iter(nums)), 0)


if __name__ == '__main__':
    unittest.main()
