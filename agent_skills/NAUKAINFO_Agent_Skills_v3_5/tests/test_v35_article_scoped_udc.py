from pathlib import Path
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import audit_frontmatter_reference_gates as afr
import udc_lookup_request as ulr


def _make_two_article_doc(path: Path) -> None:
    d = Document()
    for name in ("UDC", "AUTOR", "pip", "Назва1"):
        if name not in [s.name for s in d.styles]:
            d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    d.add_paragraph("УДК 001.1", style="UDC")
    d.add_paragraph("")
    d.add_paragraph("First Author", style="AUTOR")
    d.add_paragraph("FIRST ARTICLE TITLE", style="Назва1")
    d.add_paragraph("Body one.")

    # A DOI line with the UDC style is an intentional false-positive fixture.
    d.add_paragraph("DOI: https://doi.org/10.0000/example", style="UDC")
    d.add_paragraph("Second Author", style="AUTOR")
    d.add_paragraph("SECOND ARTICLE TITLE", style="Назва1")
    d.add_paragraph("Body two.")
    d.save(path)


def test_full_audit_catches_missing_udc_in_second_article(tmp_path):
    p = tmp_path / "two.docx"
    _make_two_article_doc(p)
    report = afr.audit(p)
    defects = [d for d in report["defects"] if d["code"] == "ARTICLE_UDC_MISSING"]
    assert len(defects) == 1
    assert defects[0]["title"] == "SECOND ARTICLE TITLE"


def test_lookup_is_article_scoped_and_doi_style_is_not_udc(tmp_path):
    p = tmp_path / "two.docx"
    _make_two_article_doc(p)
    existing = ulr.build_request(
        p, "a-1", "TEST", article_title="FIRST ARTICLE TITLE"
    )
    missing = ulr.build_request(
        p, "a-2", "TEST", article_title="SECOND ARTICLE TITLE"
    )
    assert existing["status"] == "existing_udc"
    assert missing["status"] == "UDC_LOOKUP_REQUIRED"
    assert missing["title"] == "SECOND ARTICLE TITLE"


def test_multi_article_lookup_requires_selector(tmp_path):
    p = tmp_path / "two.docx"
    _make_two_article_doc(p)
    result = ulr.build_request(p, "a-x", "TEST")
    assert result["status"] == "ARTICLE_SELECTOR_REQUIRED"
