from pathlib import Path
import tempfile
import unittest
import sys

from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
import table_format_fidelity as tff


class TableFormatFidelityTest(unittest.TestCase):
    def test_effective_style_indent_is_detected_and_repaired(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source_path = tmp / "source.docx"
            target_path = tmp / "target.docx"
            output_path = tmp / "fixed.docx"

            source = Document()
            source.styles["Normal"].paragraph_format.first_line_indent = Pt(0)
            source.add_table(rows=1, cols=1).cell(0, 0).text = "Cell text"
            source.save(source_path)

            target = Document()
            target.styles["Normal"].paragraph_format.first_line_indent = Cm(1)
            target.add_table(rows=1, cols=1).cell(0, 0).text = "Cell text"
            target.save(target_path)

            source_doc = Document(source_path)
            target_doc = Document(target_path)
            self.assertEqual(tff.effective_first_line(source_doc.tables[0].cell(0, 0).paragraphs[0])["pt"], 0.0)
            self.assertGreater(tff.effective_first_line(target_doc.tables[0].cell(0, 0).paragraphs[0])["pt"], 0.0)

            p = target_doc.tables[0].cell(0, 0).paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            target_doc.save(output_path)

            fixed_doc = Document(output_path)
            self.assertEqual(tff.effective_first_line(fixed_doc.tables[0].cell(0, 0).paragraphs[0])["pt"], 0.0)
            self.assertEqual(tff.compare_tables(source_doc, fixed_doc), [])

    def test_skill_documents_effective_inheritance_rule(self):
        text = (ROOT / "skills" / "naukainfo-table-format-fidelity" / "MODULE.md").read_text(encoding="utf-8")
        self.assertIn("effective", text.lower())
        self.assertIn('w:firstLine="0"', text)
        self.assertIn("Normal", text)
        self.assertIn("render", text.lower())


if __name__ == "__main__":
    unittest.main()
