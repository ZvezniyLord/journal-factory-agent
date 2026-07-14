from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / 'scripts'))
import normalize_captions_references as ncr


class TableFigureCaptionContractTest(unittest.TestCase):
    def test_caption_and_table_roles(self):
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / 'out.docx'
            d = Document()
            existing = {s.name for s in d.styles}
            for name in ['РИС', 'РисПід', 'TABLETEXT']:
                if name not in existing:
                    d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            p = d.add_paragraph()
            p._p.append(OxmlElement('w:drawing'))
            d.add_paragraph('Рис. 1. Тестова схема')
            d.add_paragraph('Джерело: розроблено автором.')
            d.add_paragraph('Таблиця 1')
            d.add_paragraph('Назва тестової таблиці')
            table = d.add_table(rows=1, cols=2)
            table.cell(0, 0).text = 'Ліва клітинка'
            table.cell(0, 1).text = 'Права клітинка'
            d.add_paragraph('СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ:')
            audit = {}
            ncr.normalize_table_figure_captions(d, 0, 5, audit)
            d.save(out)
            x = Document(out)
            self.assertEqual(x.paragraphs[0].style.name, 'РИС')
            self.assertEqual(x.paragraphs[1].style.name, 'РисПід')
            self.assertEqual(x.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(x.paragraphs[3].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertTrue(all(r.bold for r in x.paragraphs[3].runs if r.text))
            self.assertEqual(x.paragraphs[4].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(all(r.bold for r in x.paragraphs[4].runs if r.text))
            for cell in x.tables[0].row_cells(0):
                for cp in cell.paragraphs:
                    self.assertEqual(cp.style.name, 'TABLETEXT')
                    self.assertEqual(cp.paragraph_format.first_line_indent.pt, 0)


if __name__ == '__main__':
    unittest.main()
