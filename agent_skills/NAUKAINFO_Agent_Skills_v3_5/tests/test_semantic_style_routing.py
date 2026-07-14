from __future__ import annotations
import tempfile
import unittest
from pathlib import Path
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / 'scripts'))
import semantic_paragraph_roles as spr

class SemanticStyleRoutingTest(unittest.TestCase):
    def test_protected_roles_get_zero_first_line(self):
        with tempfile.TemporaryDirectory() as td:
            td=Path(td); src=td/'a.docx'; out=td/'b.docx'; audit=td/'a.json'
            d=Document(); d.add_paragraph('TABLE OF CONTENTS')
            for text in ['DOI: https://doi.org/10.x','УДК 001','Іванов Іван Іванович','НАЗВА НАУКОВОЇ СТАТТІ ДЛЯ ПЕРЕВІРКИ','Анотація. Текст','Ключові слова: тест','Рис. 1. Схема','Таблиця 1','Назва таблиці','СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ:','Джерело один']:
                p=d.add_paragraph(text); p.paragraph_format.first_line_indent=__import__('docx').shared.Cm(1)
            d.save(src)
            rep=spr.apply_semantic_rules(src,out,audit)
            self.assertEqual(rep['status'],'pass')
            self.assertFalse(rep['issues'])
            x=Document(out)
            by={p.text.strip():p for p in x.paragraphs if p.text.strip()}
            self.assertIsNotNone(by['Анотація. Текст'].paragraph_format.first_line_indent)
            self.assertIsNotNone(by['Ключові слова: тест'].paragraph_format.first_line_indent)

    def test_hanging_reference_is_preserved(self):
        d=Document(); p=d.add_paragraph('Reference')
        ind=OxmlElement('w:ind'); ind.set(qn('w:left'),'567'); ind.set(qn('w:hanging'),'567'); p._p.get_or_add_pPr().append(ind)
        spr.set_zero_first_line(p)
        attrs=spr.get_ind_attrs(p)
        self.assertEqual(attrs.get('hanging'),'567')
        self.assertNotIn('firstLine',attrs)

if __name__=='__main__': unittest.main()
