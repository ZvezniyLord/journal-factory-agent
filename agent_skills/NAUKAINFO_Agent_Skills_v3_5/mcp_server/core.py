from __future__ import annotations

import hashlib
import json
import os
import subprocess
import sys
import zipfile
from pathlib import Path
from typing import Any, Iterable


def _resolve(path: str | Path) -> Path:
    return Path(path).expanduser().resolve()


def _project_root(project_root: str | Path | None = None) -> Path:
    value = project_root or os.environ.get("NAUKAINFO_PROJECT_ROOT")
    if not value:
        raise ValueError("NAUKAINFO_PROJECT_ROOT or project_root is required")
    root = _resolve(value)
    if not (root / "launcher.py").exists():
        raise FileNotFoundError(f"launcher.py not found in project root: {root}")
    return root


def _python() -> str:
    return os.environ.get("NAUKAINFO_PYTHON") or sys.executable


def _run(args: list[str], cwd: Path, timeout: int = 1800) -> dict[str, Any]:
    completed = subprocess.run(
        args,
        cwd=str(cwd),
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        timeout=timeout,
        check=False,
    )
    return {
        "command": args,
        "cwd": str(cwd),
        "returncode": completed.returncode,
        "stdout": completed.stdout,
        "stderr": completed.stderr,
        "ok": completed.returncode == 0,
    }


def sha256_file(path: str | Path) -> str:
    p = _resolve(path)
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def snapshot_tree(path: str | Path, max_files: int = 20000) -> dict[str, Any]:
    root = _resolve(path)
    if root.is_file():
        return {"root": str(root), "files": {root.name: sha256_file(root)}}
    files: dict[str, str] = {}
    for i, p in enumerate(sorted(x for x in root.rglob("*") if x.is_file())):
        if i >= max_files:
            return {"root": str(root), "files": files, "truncated": True}
        files[str(p.relative_to(root))] = sha256_file(p)
    return {"root": str(root), "files": files, "truncated": False}


def _assert_output_isolated(raw_root: Path, template: Path, output: Path) -> None:
    raw_root = raw_root.resolve()
    template = template.resolve()
    output = output.resolve()
    if output == raw_root or raw_root in output.parents:
        raise ValueError("Output must not be raw-root or a child of raw-root")
    if output == template:
        raise ValueError("Output must not equal template path")
    if template.is_dir() and (output == template or template in output.parents):
        raise ValueError("Output must not be inside template directory")


def read_project_context(project_root: str | Path | None = None) -> dict[str, str]:
    root = _project_root(project_root)
    candidates = [
        root / "README_НОВА_КОНЦЕПЦІЯ.md",
        root / "docs" / "ARCHITECTURE.md",
        root / "docs" / "LLM_POLICY.md",
        root / "project_memory" / "README.md",
        root / "project_memory" / "conventions.md",
        root / "project_memory" / "decisions.md",
        root / "project_memory" / "known_bugs.md",
        root / "project_memory" / "failed_attempts.md",
        root / "project_memory" / "working_solutions.md",
        root / "project_memory" / "regression_tests.md",
        root / "project_memory" / "todo.md",
    ]
    result: dict[str, str] = {}
    for p in candidates:
        if p.exists():
            result[str(p.relative_to(root))] = p.read_text(encoding="utf-8", errors="replace")
    return result


def run_unit_tests(project_root: str | Path | None = None, timeout: int = 600) -> dict[str, Any]:
    root = _project_root(project_root)
    return _run([_python(), "-m", "unittest", "discover", "-q"], root, timeout)


def test_llm_endpoint(project_root: str | Path | None, workspace: str | Path, base_url: str, model: str, timeout_seconds: int = 300) -> dict[str, Any]:
    root = _project_root(project_root)
    workspace = _resolve(workspace)
    workspace.mkdir(parents=True, exist_ok=True)
    return _run([
        _python(), "launcher.py", "--workspace", str(workspace), "--mode", "test-llm",
        "--llm-base-url", base_url, "--llm-model", model, "--llm-timeout", str(timeout_seconds),
    ], root, timeout_seconds + 60)


def template_snapshot(project_root: str | Path | None, template: str | Path, output: str | Path) -> dict[str, Any]:
    root = _project_root(project_root)
    template_p, output_p = _resolve(template), _resolve(output)
    output_p.parent.mkdir(parents=True, exist_ok=True)
    return _run([_python(), "launcher.py", "snapshot", "--template", str(template_p), "--output", str(output_p)], root, 300)


def scan_conference(project_root: str | Path | None, raw_root: str | Path, template: str | Path, workspace: str | Path, llm_base_url: str = "", llm_model: str = "", enable_internal_llm: bool = False, timeout_seconds: int = 300) -> dict[str, Any]:
    root = _project_root(project_root)
    raw_p, template_p, workspace_p = _resolve(raw_root), _resolve(template), _resolve(workspace)
    _assert_output_isolated(raw_p, template_p, workspace_p)
    workspace_p.mkdir(parents=True, exist_ok=True)
    cmd = [
        _python(), "launcher.py", "--workspace", str(workspace_p), "--mode", "scan-conference",
        "--raw-root", str(raw_p), "--template", str(template_p), "--llm-timeout", str(timeout_seconds),
    ]
    if llm_base_url:
        cmd += ["--llm-base-url", llm_base_url]
    if llm_model:
        cmd += ["--llm-model", llm_model]
    if enable_internal_llm:
        cmd += ["--enable-llm"]
    return _run(cmd, root, timeout_seconds + 600)


def read_scan_reports(workspace: str | Path) -> dict[str, Any]:
    root = _resolve(workspace)
    names = ["scan_manifest.json", "scan_files.json", "scan_matches.json", "scan_skipped_files.json", "scan_summary.json", "template_style_snapshot.json"]
    result: dict[str, Any] = {}
    for name in names:
        p = root / name
        if p.exists():
            try:
                result[name] = json.loads(p.read_text(encoding="utf-8"))
            except Exception as exc:
                result[name] = {"_error": f"{type(exc).__name__}: {exc}"}
    return result


def inspect_docx_readonly(path: str | Path, max_paragraphs: int = 80) -> dict[str, Any]:
    from docx import Document
    p = _resolve(path)
    if p.suffix.lower() != ".docx":
        return {"path": str(p), "requires_word_conversion": True, "reason": "read-only inspector accepts DOCX only"}
    before = sha256_file(p)
    doc = Document(p)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
        paragraphs.append({"index": i, "text": para.text, "style": para.style.name if para.style else ""})
    table_cells = []
    for ti, table in enumerate(doc.tables[:10]):
        for ri, row in enumerate(table.rows[:20]):
            for ci, cell in enumerate(row.cells[:10]):
                text = " ".join(x.text.strip() for x in cell.paragraphs if x.text.strip())
                if text:
                    table_cells.append({"table": ti, "row": ri, "cell": ci, "text": text[:1000]})
    with zipfile.ZipFile(p) as z:
        names = z.namelist()
        counts = {
            "media": sum(1 for n in names if n.startswith("word/media/")),
            "embeddings": sum(1 for n in names if n.startswith("word/embeddings/")),
            "drawings": sum(1 for n in names if "drawing" in n.casefold()),
        }
    after = sha256_file(p)
    return {
        "path": str(p),
        "hash_unchanged": before == after,
        "paragraphs": paragraphs,
        "table_cells": table_cells,
        "counts": {"tables": len(doc.tables), "inline_shapes": len(doc.inline_shapes), **counts},
    }


def build_journal(project_root: str | Path | None, raw_root: str | Path, template: str | Path, output_root: str | Path, approval: str, llm_base_url: str = "", llm_model: str = "", enable_internal_llm: bool = False, enable_internal_manifest_llm: bool = False, timeout_seconds: int = 300) -> dict[str, Any]:
    if approval != "BUILD_CONFIRMED":
        raise PermissionError("Full build requires approval='BUILD_CONFIRMED'")
    root = _project_root(project_root)
    raw_p, template_p, output_p = _resolve(raw_root), _resolve(template), _resolve(output_root)
    _assert_output_isolated(raw_p, template_p, output_p)
    raw_before = snapshot_tree(raw_p)
    template_before = sha256_file(template_p)
    output_p.mkdir(parents=True, exist_ok=True)
    cmd = [
        _python(), "launcher.py", "prepare-conference", "--raw-root", str(raw_p),
        "--output-root", str(output_p), "--template", str(template_p), "--llm-timeout", str(timeout_seconds),
    ]
    if llm_base_url:
        cmd += ["--llm-base-url", llm_base_url]
    if llm_model:
        cmd += ["--llm-model", llm_model]
    if enable_internal_llm:
        cmd += ["--enable-llm"]
    if enable_internal_manifest_llm:
        cmd += ["--enable-llm-manifest"]
    result = _run(cmd, root, timeout_seconds + 3600)
    result["template_unchanged"] = template_before == sha256_file(template_p)
    result["raw_snapshot_before"] = raw_before
    # Full re-hash may be expensive but is intentional for safety.
    result["raw_unchanged"] = raw_before == snapshot_tree(raw_p)
    result["output_root"] = str(output_p)
    return result


def list_run_artifacts(run_dir: str | Path) -> dict[str, Any]:
    root = _resolve(run_dir)
    files = []
    for p in sorted(x for x in root.rglob("*") if x.is_file()):
        files.append({"path": str(p), "size": p.stat().st_size, "suffix": p.suffix.lower()})
    return {"run_dir": str(root), "files": files}


def audit_docx(project_root: str | Path | None, docx: str | Path, out_dir: str | Path) -> dict[str, Any]:
    root = _project_root(project_root)
    docx_p, out_p = _resolve(docx), _resolve(out_dir)
    out_p.mkdir(parents=True, exist_ok=True)
    return _run([_python(), "launcher.py", "audit-docx", "--docx", str(docx_p), "--out-dir", str(out_p)], root, 1800)


def quality_gate(project_root: str | Path | None, audit_dir: str | Path, output: str | Path, operator_actions: str | Path) -> dict[str, Any]:
    root = _project_root(project_root)
    return _run([
        _python(), "launcher.py", "quality-gate", "--audit-dir", str(_resolve(audit_dir)),
        "--output", str(_resolve(output)), "--operator-actions", str(_resolve(operator_actions)),
    ], root, 600)


def render_docx_pdf(docx: str | Path, output_pdf: str | Path) -> dict[str, Any]:
    p, out = _resolve(docx), _resolve(output_pdf)
    out.parent.mkdir(parents=True, exist_ok=True)
    before = sha256_file(p)
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        return {"ok": False, "error": f"pywin32/Word COM unavailable: {type(exc).__name__}: {exc}"}
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = None
    try:
        doc = word.Documents.Open(str(p), ReadOnly=True, AddToRecentFiles=False)
        pages = int(doc.ComputeStatistics(2))  # wdStatisticPages
        doc.ExportAsFixedFormat(str(out), 17)  # wdExportFormatPDF
        return {"ok": True, "pdf": str(out), "pages": pages, "source_hash_unchanged": before == sha256_file(p)}
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()
