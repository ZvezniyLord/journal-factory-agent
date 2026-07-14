from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SLOT = '[[NAUKAINFO_ARTICLE_SLOT]]'
END = '[[NAUKAINFO_MASTER_END]]'


def _paragraph_text(el) -> str:
    return ''.join(el.xpath('.//w:t/text()'))


def _make_marker(text: str):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def prepare_template(template_path: Path, prepared_path: Path) -> None:
    doc = Document(template_path)
    body = doc._element.body

    # The insertion point is immediately before the empty paragraph that owns
    # the section break after TABLE OF CONTENTS. This keeps the protected tail
    # page intact and makes inserted content start after the existing page break.
    section_break_paragraph = None
    seen_toc = False
    for p in body.findall(qn('w:p')):
        txt = _paragraph_text(p).strip()
        if txt == 'TABLE OF CONTENTS':
            seen_toc = True
            continue
        if seen_toc:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                section_break_paragraph = p
                break
    if section_break_paragraph is None:
        raise RuntimeError('Could not find the section-break paragraph after TABLE OF CONTENTS')

    section_break_paragraph.addprevious(_make_marker(SLOT))

    # Add an end marker immediately before the document-level sectPr.
    body_sectpr = body.find(qn('w:sectPr'))
    end_marker = _make_marker(END)
    if body_sectpr is not None:
        body_sectpr.addprevious(end_marker)
    else:
        body.append(end_marker)

    doc.save(prepared_path)


def compose_and_relocate(prepared_template: Path, article_path: Path, output_path: Path) -> None:
    master = Document(prepared_template)
    composer = Composer(master)
    composer.append(Document(article_path))
    composed_temp = output_path.with_suffix('.composed.tmp.docx')
    composer.save(composed_temp)

    doc = Document(composed_temp)
    body = doc._element.body
    children = list(body)

    slot_el = next((el for el in children if el.tag == qn('w:p') and SLOT in _paragraph_text(el)), None)
    end_el = next((el for el in children if el.tag == qn('w:p') and END in _paragraph_text(el)), None)
    if slot_el is None or end_el is None:
        raise RuntimeError('Template markers were not preserved during composition')

    end_idx = list(body).index(end_el)
    body_sectpr = body.find(qn('w:sectPr'))
    stop_idx = list(body).index(body_sectpr) if body_sectpr is not None else len(body)

    appended = list(body)[end_idx + 1:stop_idx]
    if not appended:
        raise RuntimeError('No appended article content detected')

    # Move article elements, in order, to the slot before the protected tail section.
    for el in appended:
        slot_el.addprevious(el)

    body.remove(slot_el)
    body.remove(end_el)

    doc.save(output_path)
    composed_temp.unlink(missing_ok=True)


def build(template: str, article: str, output: str) -> None:
    template_path = Path(template)
    article_path = Path(article)
    output_path = Path(output)
    prepared = output_path.with_suffix('.prepared.tmp.docx')
    prepare_template(template_path, prepared)
    compose_and_relocate(prepared, article_path, output_path)
    prepared.unlink(missing_ok=True)


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('template')
    ap.add_argument('article')
    ap.add_argument('output')
    args = ap.parse_args()
    build(args.template, args.article, args.output)
