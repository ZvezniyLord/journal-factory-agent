#!/usr/bin/env python3
"""Conservative DOCX source/final article-body integrity audit.

The audit deliberately fails closed. It compares visible paragraph/table text and
manual-vs-automatic list mode. It does not rewrite documents.
"""
from __future__ import annotations
import argparse
import json
import re
from dataclasses import dataclass, asdict
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any
from docx import Document
from docx.oxml.ns import qn

REF_HEADINGS = {
    "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ:",
    "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ",
    "REFERENCES",
    "REFERENCES:",
}
ANNOTATION_RE = re.compile(r"^(Анотація[\.:]?|Abstract[\.:]?)\s*", re.I)
KEYWORDS_RE = re.compile(r"^(Ключові\s+слова:|Keywords:)\s*", re.I)


def norm_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u00a0", " ")).strip()


def norm_allowed_labels(text: str) -> str:
    text = norm_ws(text)
    text = ANNOTATION_RE.sub("ANNOTATION_LABEL ", text)
    text = KEYWORDS_RE.sub("KEYWORDS_LABEL ", text)
    return text.strip()


def list_mode(paragraph) -> str:
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.find(qn("w:numPr")) is not None:
        return "automatic"
    t = norm_ws(paragraph.text)
    if re.match(r"^(?:\(?\d+[\.)]|[-–—•▪◦]|[A-Za-zА-Яа-яІіЇїЄє]\))\s+", t):
        return "manual"
    return "none"


def body_start_index(paragraphs) -> int:
    # Prefer first annotation/abstract. Otherwise use first paragraph after an article-title style.
    for i, p in enumerate(paragraphs):
        if ANNOTATION_RE.match(norm_ws(p.text)):
            return i
    for i, p in enumerate(paragraphs):
        if getattr(p.style, "style_id", "") in {"1", "Назва1", "Title"}:
            j = i + 1
            while j < len(paragraphs) and not norm_ws(paragraphs[j].text):
                j += 1
            return j
    return 0


def extract(doc_path: Path) -> dict[str, Any]:
    doc = Document(str(doc_path))
    paragraphs = list(doc.paragraphs)
    start = body_start_index(paragraphs)
    p_items = []
    for idx, p in enumerate(paragraphs[start:], start=start):
        txt = norm_ws(p.text)
        p_items.append({
            "index": idx,
            "text": txt,
            "normalized": norm_allowed_labels(txt),
            "list_mode": list_mode(p),
            "style_id": getattr(p.style, "style_id", ""),
        })
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([norm_ws(cell.text) for cell in row.cells])
        tables.append({"index": ti, "rows": rows})
    xml = doc.element.body.xml
    objects = {
        "drawings": xml.count("<w:drawing"),
        "pict": xml.count("<w:pict"),
        "objects": xml.count("<w:object"),
        "math": xml.count("<m:oMath") + xml.count("<m:oMathPara"),
    }
    return {"paragraphs": p_items, "tables": tables, "objects": objects}


def similarity(a: list[str], b: list[str]) -> float:
    return SequenceMatcher(a=a, b=b, autojunk=False).ratio()


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("source")
    ap.add_argument("final")
    ap.add_argument("--out-json")
    args = ap.parse_args()
    src = extract(Path(args.source))
    dst = extract(Path(args.final))

    src_text = [p["normalized"] for p in src["paragraphs"]]
    dst_text = [p["normalized"] for p in dst["paragraphs"]]
    src_nonempty = [x for x in src_text if x]
    dst_nonempty = [x for x in dst_text if x]

    text_ratio = similarity(src_nonempty, dst_nonempty)
    structure_ratio = similarity(
        [f'{p["list_mode"]}:{p["normalized"]}' for p in src["paragraphs"]],
        [f'{p["list_mode"]}:{p["normalized"]}' for p in dst["paragraphs"]],
    )
    tables_equal = src["tables"] == dst["tables"]
    objects_equal = src["objects"] == dst["objects"]

    report = {
        "source": args.source,
        "final": args.final,
        "body_text_similarity": text_ratio,
        "body_structure_similarity": structure_ratio,
        "tables_exact": tables_equal,
        "objects_exact": objects_equal,
        "source_objects": src["objects"],
        "final_objects": dst["objects"],
        "pass": text_ratio == 1.0 and structure_ratio >= 0.99 and tables_equal and objects_equal,
        "policy": {
            "lexical_target": 1.0,
            "structural_minimum": 0.99,
            "manual_body_lists_must_remain_manual": True,
            "reference_numbering_exception": True,
        },
    }
    payload = json.dumps(report, ensure_ascii=False, indent=2)
    if args.out_json:
        Path(args.out_json).write_text(payload, encoding="utf-8")
    print(payload)
    return 0 if report["pass"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
