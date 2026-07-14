from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


def length_pt(value):
    return None if value is None else round(value.pt, 4)


def effective_first_line(paragraph) -> dict[str, Any]:
    direct = paragraph.paragraph_format.first_line_indent
    if direct is not None:
        return {"source": "direct", "pt": round(direct.pt, 4)}
    style = paragraph.style
    seen: set[str] = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        value = style.paragraph_format.first_line_indent
        if value is not None:
            return {"source": f"style:{style.name}", "pt": round(value.pt, 4)}
        style = style.base_style
    return {"source": "default", "pt": 0.0}


def run_signature(paragraph):
    return [
        {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": str(run.underline),
            "size_pt": length_pt(run.font.size),
            "font": run.font.name,
        }
        for run in paragraph.runs
    ]


def paragraph_signature(paragraph):
    pf = paragraph.paragraph_format
    return {
        "text": paragraph.text,
        "alignment": None if paragraph.alignment is None else int(paragraph.alignment),
        "effective_first_line": effective_first_line(paragraph),
        "left_indent_pt": length_pt(pf.left_indent),
        "right_indent_pt": length_pt(pf.right_indent),
        "space_before_pt": length_pt(pf.space_before),
        "space_after_pt": length_pt(pf.space_after),
        "line_spacing": pf.line_spacing,
        "keep_together": pf.keep_together,
        "keep_with_next": pf.keep_with_next,
        "page_break_before": pf.page_break_before,
        "widow_control": pf.widow_control,
        "runs": run_signature(paragraph),
    }


def compare_tables(source_doc, target_doc) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    if len(source_doc.tables) != len(target_doc.tables):
        return [{"kind": "table_count", "source": len(source_doc.tables), "target": len(target_doc.tables)}]

    for ti, (source_table, target_table) in enumerate(zip(source_doc.tables, target_doc.tables), start=1):
        source_shape = (len(source_table.rows), len(source_table.columns))
        target_shape = (len(target_table.rows), len(target_table.columns))
        if source_shape != target_shape:
            issues.append({"kind": "table_shape", "table": ti, "source": source_shape, "target": target_shape})
            continue

        for ri, (source_row, target_row) in enumerate(zip(source_table.rows, target_table.rows), start=1):
            for ci, (source_cell, target_cell) in enumerate(zip(source_row.cells, target_row.cells), start=1):
                if source_cell.text != target_cell.text:
                    issues.append({"kind": "cell_text", "table": ti, "row": ri, "cell": ci})
                if source_cell.vertical_alignment != target_cell.vertical_alignment:
                    issues.append({
                        "kind": "vertical_alignment", "table": ti, "row": ri, "cell": ci,
                        "source": str(source_cell.vertical_alignment), "target": str(target_cell.vertical_alignment),
                    })
                if len(source_cell.paragraphs) != len(target_cell.paragraphs):
                    issues.append({
                        "kind": "paragraph_count", "table": ti, "row": ri, "cell": ci,
                        "source": len(source_cell.paragraphs), "target": len(target_cell.paragraphs),
                    })
                    continue

                for pi, (source_p, target_p) in enumerate(zip(source_cell.paragraphs, target_cell.paragraphs), start=1):
                    ss = paragraph_signature(source_p)
                    ts = paragraph_signature(target_p)
                    for key in [
                        "text", "alignment", "left_indent_pt", "right_indent_pt", "space_before_pt",
                        "space_after_pt", "line_spacing", "keep_together", "keep_with_next",
                        "page_break_before", "widow_control", "runs",
                    ]:
                        if ss[key] != ts[key]:
                            issues.append({
                                "kind": f"paragraph_{key}", "table": ti, "row": ri, "cell": ci,
                                "paragraph": pi, "source": ss[key], "target": ts[key],
                            })
                    if ss["effective_first_line"]["pt"] != ts["effective_first_line"]["pt"]:
                        issues.append({
                            "kind": "effective_first_line_indent", "table": ti, "row": ri, "cell": ci,
                            "paragraph": pi, "source": ss["effective_first_line"],
                            "target": ts["effective_first_line"],
                        })
    return issues


def main() -> None:
    parser = argparse.ArgumentParser(description="Audit and repair table paragraph formatting inherited from a target template.")
    parser.add_argument("source", type=Path, help="Original/normalized article used as table-format reference")
    parser.add_argument("target", type=Path, help="Journal/template copy to audit")
    parser.add_argument("output", type=Path, help="Repaired DOCX output path")
    parser.add_argument("--audit-json", type=Path, required=True)
    parser.add_argument("--audit-only", action="store_true")
    args = parser.parse_args()

    source_doc = Document(args.source)
    target_doc = Document(args.target)
    before = compare_tables(source_doc, target_doc)
    structural = [issue for issue in before if issue["kind"] in {"table_count", "table_shape", "cell_text", "paragraph_count"}]
    if structural and not args.audit_only:
        raise SystemExit("Refusing automatic repair because table structure/text differs; inspect audit output.")

    fixes: list[dict[str, Any]] = []
    if not args.audit_only:
        for ti, (source_table, target_table) in enumerate(zip(source_doc.tables, target_doc.tables), start=1):
            for ri, (source_row, target_row) in enumerate(zip(source_table.rows, target_table.rows), start=1):
                for ci, (source_cell, target_cell) in enumerate(zip(source_row.cells, target_row.cells), start=1):
                    for pi, (source_p, target_p) in enumerate(zip(source_cell.paragraphs, target_cell.paragraphs), start=1):
                        source_indent = effective_first_line(source_p)
                        target_indent = effective_first_line(target_p)
                        if source_indent["pt"] != target_indent["pt"]:
                            target_p.paragraph_format.first_line_indent = Pt(source_indent["pt"])
                            fixes.append({
                                "table": ti, "row": ri, "cell": ci, "paragraph": pi,
                                "text_preview": target_p.text[:80],
                                "before": target_indent,
                                "source_reference": source_indent,
                                "after": effective_first_line(target_p),
                            })
        target_doc.save(args.output)
        reopened = Document(args.output)
        after = compare_tables(source_doc, reopened)
    else:
        after = before

    report = {
        "source": str(args.source),
        "target": str(args.target),
        "output": None if args.audit_only else str(args.output),
        "before_issue_count": len(before),
        "before_issues": before,
        "fix_count": len(fixes),
        "fixes": fixes,
        "after_issue_count": len(after),
        "after_issues": after,
        "status": "pass" if not after else ("needs_operator_review" if not structural else "fail"),
    }
    args.audit_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({k: report[k] for k in ["before_issue_count", "fix_count", "after_issue_count", "status"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
