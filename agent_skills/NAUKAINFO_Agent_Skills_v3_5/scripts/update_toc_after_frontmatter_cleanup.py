import re, json, fitz
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX=Path('/mnt/data/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v26_SCOPE_HEADERS.docx')
PDF=Path('/mnt/data/render_full_v26/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v26_SCOPE_HEADERS.pdf')
OUT=Path('/mnt/data/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v26_FINAL.docx')
QA=Path('/mnt/data/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v26_FINAL_QA.json')

def norm(s):
    s=s.upper().replace('–','-').replace('—','-')
    s=re.sub(r'\s+',' ',s).strip()
    return s

def set_cell(cell, text, style_id):
    # clear all content in cell
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p=cell.add_paragraph()
    try: p.style=style_id
    except Exception: pass
    p.add_run(text)
    return p

def get_page_map(titles):
    d=fitz.open(str(PDF)); texts=[p.get_text('text') for p in d]
    pages=[]
    for t in titles:
        nt=norm(t); found=None
        for i,txt in enumerate(texts[6:], start=7):
            if nt and nt in norm(txt): found=i; break
        if found is None:
            key=re.sub(r'[^A-ZА-ЯІЇЄҐ0-9]+','',nt)[:45]
            for i,txt in enumerate(texts[6:], start=7):
                txtkey=re.sub(r'[^A-ZА-ЯІЇЄҐ0-9]+','',norm(txt))
                if key and key in txtkey: found=i; break
        pages.append(found)
    return pages

def article_records(doc):
    records=[]; current_section=None
    for i,p in enumerate(doc.paragraphs):
        if p.style.style_id=='SECTION':
            current_section=p.text.strip()
        if p.style.style_id=='11':
            # collect AUTOR lines between previous UDC/SECTION and title
            authors=[]
            j=i-1
            while j>=0:
                pj=doc.paragraphs[j]
                if pj.style.style_id in ('SECTION','UDC'):
                    break
                if pj.style.style_id=='AUTOR' and pj.text.strip():
                    authors.insert(0, re.sub(r'\s+',' ',pj.text.strip()))
                j-=1
            records.append({'section':current_section,'authors':', '.join(authors),'title':re.sub(r'\s+',' ',p.text.strip()),'para_index':i})
    pages=get_page_map([r['title'] for r in records])
    for r,p in zip(records,pages): r['page']=p
    return records

if __name__=='__main__':
    doc=Document(str(DOCX))
    records=article_records(doc)
    tbl=doc.tables[0]
    # Rebuild existing rows according to records. No row creation needed: 14 sections + 24 article pairs already exists.
    row_idx=0; art_idx=0; article_num=1; last_section=None
    for r in records:
        if r['section'] != last_section:
            # set merged section row
            row=tbl.rows[row_idx]
            seen=set()
            for c in row.cells:
                tc_id=id(c._tc)
                if tc_id not in seen:
                    set_cell(c, r['section'], 'TabSEC')
                    seen.add(tc_id)
            row_idx+=1
            last_section=r['section']
        row=tbl.rows[row_idx]
        set_cell(row.cells[0], f'{article_num}.', 'TabPIP')
        set_cell(row.cells[1], r['authors'], 'TabPIP')
        set_cell(row.cells[2], str(r['page'] or ''), 'TabPIP')
        row_idx+=1
        row=tbl.rows[row_idx]
        set_cell(row.cells[0], '', 'TabTaitl')
        set_cell(row.cells[1], r['title'], 'TabTaitl')
        set_cell(row.cells[2], '', 'TabTaitl')
        row_idx+=1
        article_num+=1
    doc.save(str(OUT))
    qa={'input':str(DOCX),'output':str(OUT),'article_count':len(records),'page_count_pdf':len(fitz.open(str(PDF))),'records':records,'toc_rows_used':row_idx}
    QA.write_text(json.dumps(qa,ensure_ascii=False,indent=2),encoding='utf-8')
    print(json.dumps({'out':str(OUT),'records':len(records),'toc_rows_used':row_idx},ensure_ascii=False,indent=2))
