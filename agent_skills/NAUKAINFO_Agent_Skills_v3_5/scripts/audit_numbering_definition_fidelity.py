#!/usr/bin/env python3
"""Audit/fix DOCX numbering definition fidelity for known NAUKAINFO merge regressions.

This script is intentionally conservative. It never changes body text. It only checks or repairs
numbering.xml mappings for known body-list paragraphs whose source marker is a bullet.
"""
from __future__ import annotations
from pathlib import Path
import argparse, json, zipfile
from lxml import etree
from docx import Document

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': W}
def q(tag: str) -> str:
    return f'{{{W}}}{tag}'

HNYSIUK_TARGET_PREFIXES = [
    'Пряма форма включає заробітну плату',
    'Непряма матеріальна мотивація спрямована',
]

def read_numbering(docx_path: Path):
    with zipfile.ZipFile(docx_path) as z:
        return etree.fromstring(z.read('word/numbering.xml'))

def paragraph_numid(p):
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None or pPr.numPr.numId is None:
        return None
    return str(pPr.numPr.numId.val)

def number_format(root, numid: str):
    numel = root.xpath(f'.//w:num[@w:numId="{numid}"]', namespaces=NS)
    if not numel:
        return None
    absid = numel[0].xpath('./w:abstractNumId/@w:val', namespaces=NS)[0]
    absel = root.xpath(f'.//w:abstractNum[@w:abstractNumId="{absid}"]', namespaces=NS)[0]
    fmt = absel.xpath('./w:lvl[@w:ilvl="0"]/w:numFmt/@w:val', namespaces=NS)
    txt = absel.xpath('./w:lvl[@w:ilvl="0"]/w:lvlText/@w:val', namespaces=NS)
    return {'numId': numid, 'abstractNumId': absid, 'numFmt': fmt[0] if fmt else None, 'lvlText': txt[0] if txt else None}

def find_bullet_abstract(root):
    for ab in root.xpath('.//w:abstractNum', namespaces=NS):
        fmt = ab.xpath('./w:lvl[@w:ilvl="0"]/w:numFmt/@w:val', namespaces=NS)
        txt = ab.xpath('./w:lvl[@w:ilvl="0"]/w:lvlText/@w:val', namespaces=NS)
        if fmt == ['bullet'] and txt == ['-']:
            return ab.get(q('abstractNumId'))
    return None

def audit(docx_path: Path):
    doc = Document(docx_path)
    root = read_numbering(docx_path)
    rows = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if any(t.startswith(prefix) for prefix in HNYSIUK_TARGET_PREFIXES):
            n = paragraph_numid(p)
            rows.append({'paragraph_index': i, 'text': t[:160], 'num': number_format(root, n) if n else None})
    ok = len(rows) == 2 and all(r['num'] and r['num']['numFmt'] == 'bullet' and r['num']['lvlText'] == '-' for r in rows)
    return {'path': str(docx_path), 'target_paragraphs': rows, 'pass': ok}

def fix(in_path: Path, out_path: Path):
    root = read_numbering(in_path)
    bullet_abs = find_bullet_abstract(root)
    if bullet_abs is None:
        raise RuntimeError('Cannot find bullet abstractNum with lvlText -')
    changed = []
    doc = Document(in_path)
    target_numids = set()
    for p in doc.paragraphs:
        t = p.text.strip()
        if any(t.startswith(prefix) for prefix in HNYSIUK_TARGET_PREFIXES):
            n = paragraph_numid(p)
            if n:
                target_numids.add(n)
    for n in sorted(target_numids):
        nf = number_format(root, n)
        if nf and nf['numFmt'] != 'bullet':
            numel = root.xpath(f'.//w:num[@w:numId="{n}"]', namespaces=NS)[0]
            abs_el = numel.xpath('./w:abstractNumId', namespaces=NS)[0]
            before = abs_el.get(q('val'))
            abs_el.set(q('val'), bullet_abs)
            changed.append({'numId': n, 'from': before, 'to': bullet_abs})
    new_numbering = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
    with zipfile.ZipFile(in_path, 'r') as zin, zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/numbering.xml':
                data = new_numbering
            zout.writestr(item, data)
    return changed

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--fix-out')
    ap.add_argument('--out-json')
    args = ap.parse_args()
    path = Path(args.docx)
    result = audit(path)
    if args.fix_out and not result['pass']:
        changed = fix(path, Path(args.fix_out))
        result = audit(Path(args.fix_out))
        result['changed'] = changed
    if args.out_json:
        Path(args.out_json).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding='utf-8')
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(0 if result['pass'] else 2)

if __name__ == '__main__':
    main()
