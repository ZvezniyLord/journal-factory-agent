from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOI_RE = re.compile(r'^\s*(?:DOI\s*:?|https?://doi\.org/)', re.I)
UDC_RE = re.compile(r'^\s*(?:УДК|UDC)\b', re.I)
FIG_CAP_RE = re.compile(r'^\s*(?:Рис(?:унок)?\.?|Fig(?:ure)?\.?)\s*\d+', re.I)
TABLE_CAP_RE = re.compile(r'^\s*(?:Таблиця|Table)\s*\d+', re.I)
SOURCE_RE = re.compile(r'^\s*(?:Джерело|Source)\s*:', re.I)
REF_HEAD_RE = re.compile(r'^\s*(?:СПИСОК\s+(?:ВИКОРИСТАНИХ\s+ДЖЕРЕЛ|ВИКОРИСТАНОЇ\s+ЛІТЕРАТУРИ|ДЖЕРЕЛ)|REFERENCES|BIBLIOGRAPHY)\s*:?\s*$', re.I)
ABSTRACT_RE = re.compile(r'^\s*(?:Анотац(?:ія|iя)|Abstract)\s*[.:]', re.I)
KEYWORDS_RE = re.compile(r'^\s*(?:Ключові\s+слова|Keywords)\s*:', re.I)
TOC_RE = re.compile(r'^\s*TABLE\s+OF\s+CONTENTS\s*$', re.I)
TAIL_RE = re.compile(r'^\s*SCIENCE\s+IN\s+THE\s+MODERN\s+WORLD\s*$', re.I)


def iter_block_items(parent: _Document | _Cell):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def has_drawing(p: Paragraph) -> bool:
    return bool(p._p.xpath('.//*[local-name()="drawing" or local-name()="pict" or local-name()="AlternateContent"]'))


def has_numpr(p: Paragraph) -> bool:
    ppr = p._p.pPr
    return bool(ppr is not None and ppr.find(qn('w:numPr')) is not None)


def get_ind_attrs(p: Paragraph) -> dict[str, str]:
    ppr = p._p.pPr
    if ppr is None:
        return {}
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        return {}
    return {k.split('}')[-1]: v for k, v in ind.attrib.items()}


def set_zero_first_line(p: Paragraph) -> bool:
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        ppr.append(ind)
    # Hanging indent already keeps the first line at the left edge; do not add a
    # conflicting firstLine attribute. Remove any positive firstLine instead.
    if ind.get(qn('w:hanging')) is not None:
        if ind.get(qn('w:firstLine')) is not None:
            del ind.attrib[qn('w:firstLine')]
            return True
        return False
    old = ind.get(qn('w:firstLine'))
    if old != '0':
        ind.set(qn('w:firstLine'), '0')
        return True
    return False


def uppercase_ratio(text: str) -> float:
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return 0.0
    return sum(c.isupper() for c in letters) / len(letters)


def likely_article_title(p: Paragraph) -> bool:
    t = ' '.join(p.text.split())
    if len(t) < 25 or DOI_RE.match(t) or UDC_RE.match(t):
        return False
    if p.style and p.style.name == 'Назва1':
        return True
    return uppercase_ratio(t) >= 0.78 and (p.alignment is None or int(p.alignment) in (1, 3))


def role_style_suggestion(role: str) -> str | None:
    return {
        'udc': 'UDC',
        'article_title': 'Назва1',
        'figure_anchor': 'РИС',
        'figure_caption': 'РисПід',
        'reference_heading': 'REF-TITLE',
        'reference_entry': 'REFER',
        'table_cell': 'TABLETEXT',
        'author_name': 'AUTOR',
        'author_detail': 'pip',
    }.get(role)


def apply_style_if_safe(p: Paragraph, role: str, doc: _Document, enabled: bool) -> str | None:
    suggestion = role_style_suggestion(role)
    if not enabled or not suggestion:
        return None
    try:
        doc.styles[suggestion]
    except KeyError:
        return None
    # Conservative assignment. Author and reference styles may affect outline
    # levels or numbering, so only report them unless explicitly requested by a
    # future, separately validated pipeline stage.
    if role in {'author_name', 'author_detail', 'reference_entry'}:
        return None
    p.style = suggestion
    return suggestion


def classify_paragraphs(doc: _Document):
    paragraphs = doc.paragraphs
    toc_idx = next((i for i, p in enumerate(paragraphs) if TOC_RE.match(' '.join(p.text.split()))), -1)
    start = toc_idx + 1 if toc_idx >= 0 else 0
    tail_candidates = [i for i, p in enumerate(paragraphs) if i > start and TAIL_RE.match(' '.join(p.text.split()))]
    end = tail_candidates[-1] if tail_candidates else len(paragraphs)

    roles: dict[int, str] = {}
    in_refs = False
    awaiting_title = False
    metadata_mode = False
    last_table_caption_idx: int | None = None
    seen_title = False

    for i in range(start, end):
        p = paragraphs[i]
        t = ' '.join(p.text.split())
        if not t and not has_drawing(p):
            continue

        if has_drawing(p):
            roles[i] = 'figure_anchor'
            continue
        if DOI_RE.match(t):
            roles[i] = 'doi'
            metadata_mode = True
            awaiting_title = True
            in_refs = False
            continue
        if UDC_RE.match(t):
            roles[i] = 'udc'
            metadata_mode = True
            awaiting_title = True
            in_refs = False
            continue
        if likely_article_title(p) and awaiting_title:
            roles[i] = 'article_title'
            metadata_mode = False
            awaiting_title = False
            seen_title = True
            continue
        if metadata_mode and t:
            # First line after UDC/DOI is the author name; subsequent lines are details.
            prior_meta = [r for j, r in roles.items() if start <= j < i and r in {'author_name', 'author_detail'}]
            roles[i] = 'author_name' if not prior_meta or (p.alignment is not None and int(p.alignment) == 2 and t.count(' ') in (2, 3) and not any(ch in t for ch in ',;:/')) else 'author_detail'
            continue
        if ABSTRACT_RE.match(t):
            roles[i] = 'abstract'
            continue
        if KEYWORDS_RE.match(t):
            roles[i] = 'keywords'
            continue
        if FIG_CAP_RE.match(t):
            roles[i] = 'figure_caption'
            continue
        if TABLE_CAP_RE.match(t):
            roles[i] = 'table_caption'
            last_table_caption_idx = i
            continue
        if SOURCE_RE.match(t):
            roles[i] = 'source_note'
            continue
        if REF_HEAD_RE.match(t):
            roles[i] = 'reference_heading'
            in_refs = True
            continue
        if in_refs:
            # A new article starts when DOI/UDC appears; handled above.
            roles[i] = 'reference_entry'
            continue
        if has_numpr(p):
            roles[i] = 'list_item'
            continue
        if last_table_caption_idx is not None and i == last_table_caption_idx + 1 and len(t) < 240:
            roles[i] = 'table_title'
            last_table_caption_idx = None
            continue

    return roles, start, end


def process_table(table: Table, changes: list[dict], table_index: int):
    seen = set()
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            for p_idx, p in enumerate(cell.paragraphs):
                before = get_ind_attrs(p)
                changed = set_zero_first_line(p)
                if changed:
                    changes.append({'role': 'table_cell', 'table': table_index, 'row': r_idx, 'cell': c_idx, 'paragraph': p_idx, 'before': before, 'after': get_ind_attrs(p), 'text': p.text[:160]})
            for nested in cell.tables:
                process_table(nested, changes, table_index)


def apply_semantic_rules(src: Path, out: Path, audit_json: Path, apply_styles: bool = False):
    doc = Document(src)
    roles, start, end = classify_paragraphs(doc)
    changes: list[dict] = []
    protected_roles = {
        'doi', 'udc', 'author_name', 'author_detail', 'article_title',
        'figure_anchor', 'figure_caption', 'table_caption', 'table_title', 'source_note',
        'reference_heading', 'reference_entry', 'list_item'
    }
    style_assignments = []

    for idx, role in sorted(roles.items()):
        p = doc.paragraphs[idx]
        before = get_ind_attrs(p)
        changed = False
        if role in protected_roles:
            changed = set_zero_first_line(p)
        assigned = apply_style_if_safe(p, role, doc, apply_styles)
        if assigned:
            style_assignments.append({'paragraph': idx, 'role': role, 'style': assigned})
        if changed:
            changes.append({'role': role, 'paragraph': idx, 'before': before, 'after': get_ind_attrs(p), 'text': p.text[:200]})

    for t_idx, table in enumerate(doc.tables):
        process_table(table, changes, t_idx)

    doc.save(out)

    # Reopen and verify no protected role has a positive direct first-line indent.
    reopened = Document(out)
    roles2, start2, end2 = classify_paragraphs(reopened)
    issues = []
    for idx, role in sorted(roles2.items()):
        if role not in protected_roles:
            continue
        attrs = get_ind_attrs(reopened.paragraphs[idx])
        val = attrs.get('firstLine')
        if val is not None and int(val) > 0:
            issues.append({'paragraph': idx, 'role': role, 'firstLine': val, 'text': reopened.paragraphs[idx].text[:160]})
    for t_idx, table in enumerate(reopened.tables):
        seen = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                for p_idx, p in enumerate(cell.paragraphs):
                    attrs = get_ind_attrs(p)
                    val = attrs.get('firstLine')
                    if val is not None and int(val) > 0:
                        issues.append({'role': 'table_cell', 'table': t_idx, 'row': r_idx, 'cell': c_idx, 'paragraph': p_idx, 'firstLine': val, 'text': p.text[:160]})

    report = {
        'source': str(src), 'output': str(out), 'article_range': [start, end],
        'roles': [{'paragraph': i, 'role': r, 'suggested_style': role_style_suggestion(r), 'text': doc.paragraphs[i].text[:200]} for i, r in sorted(roles.items())],
        'changes': changes, 'style_assignments': style_assignments,
        'table_count': len(doc.tables), 'issues': issues,
        'status': 'pass' if not issues else 'fail'
    }
    audit_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
    return report


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('src', type=Path)
    ap.add_argument('out', type=Path)
    ap.add_argument('--audit-json', type=Path, required=True)
    ap.add_argument('--apply-safe-styles', action='store_true')
    a = ap.parse_args()
    report = apply_semantic_rules(a.src, a.out, a.audit_json, a.apply_safe_styles)
    print(json.dumps({'status': report['status'], 'changes': len(report['changes']), 'issues': len(report['issues'])}, ensure_ascii=False))

if __name__ == '__main__':
    main()
