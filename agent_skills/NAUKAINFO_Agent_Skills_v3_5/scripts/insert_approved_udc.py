#!/usr/bin/env python3
"""Insert an operator-approved UDC into a DOCX and enforce one blank paragraph after it."""
from __future__ import annotations
import argparse, shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text: p.add_run(text)
    if style: p.style = style
    return p

def process(src:Path,dst:Path,udc:str):
    shutil.copy2(src,dst); doc=Document(dst)
    title = next((p for p in doc.paragraphs if p.text.strip() and p.text.strip().upper()==p.text.strip() and len(p.text.strip())>20),None)
    if title is None: raise RuntimeError('article title not found')
    p=insert_paragraph_before(title,udc,'UDC')
    blank=insert_paragraph_before(title,"",'Normal')
    doc.save(dst)

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('src',type=Path); ap.add_argument('dst',type=Path); ap.add_argument('--udc',required=True)
    a=ap.parse_args(); process(a.src,a.dst,a.udc)
if __name__=='__main__': main()
