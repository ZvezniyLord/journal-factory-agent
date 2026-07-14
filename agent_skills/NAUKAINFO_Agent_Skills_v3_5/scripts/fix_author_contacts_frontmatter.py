import re, json, shutil, zipfile
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

IN=Path('/mnt/data/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v25.docx')
OUT=Path('/mnt/data/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v26_SCOPE_HEADERS.docx')
QA=Path('/mnt/data/JOURNAL_136_FULL_24_ARTICLES_RELEASE_v26_QA.json')

EMAIL_RE=re.compile(r'(?<![\w.-])([A-Za-z0-9._%+\-]+)\s*@\s*([A-Za-z0-9.\-]+\.[A-Za-z]{2,})(?![\w.-])')
PERSONAL_EMAIL_EXCEPTIONS={'journal@naukainfo.com'}
PHONE_RE=re.compile(r'(?:\+?\d[\d\s()\-]{7,}\d)')

def norm_email(s):
    m=EMAIL_RE.search(s)
    if not m: return None
    return (m.group(1)+'@'+m.group(2)).lower()

def is_personal_email_line(text):
    e=norm_email(text)
    if not e: return False
    if e in PERSONAL_EMAIL_EXCEPTIONS: return False
    # if line mostly email or frontmatter contact line, remove whole line
    return True

def is_phone_line(text):
    # Author phone lines are not published; avoid touching DOI/UDC pages by only using this in frontmatter
    return bool(PHONE_RE.search(text))

def is_section_note(text):
    t=text.strip()
    return bool(re.match(r'^(Секц(?:ія|iя)|Section)\b', t, flags=re.I))

def clean_text(text):
    text=text.replace('\xa0',' ')
    text=re.sub(r'\s+',' ',text).strip()
    return text

def clean_author_name(text):
    t=clean_text(text)
    t=t.rstrip(' ,;')
    return t

def looks_like_person_name(t):
    t=clean_text(t).strip('.,;')
    if not t or len(t)>80: return False
    # remove initials punctuation for counting
    words=t.split()
    if len(words)<2 or len(words)>4: return False
    bad=['university','department','academy','кафедр','університет','міжнародний','національний','старший','senior','lecturer','професор','доцент','здобувач','аспірант','м.','orcid','theatre','cinema','television','artist','artists','ukraine','україна','member','union','департамент','відділ']
    low=t.lower()
    if any(b in low for b in bad): return False
    # Capitalized/initial-like tokens
    ok=0
    for w in words:
        ww=w.strip('.,;:')
        if re.match(r'^[А-ЯІЇЄҐA-Z][а-яіїєґa-zA-ZА-ЯІЇЄҐ\-’\']+$', ww) or re.match(r'^[А-ЯІЇЄҐA-Z]\.?$', ww):
            ok+=1
    return ok==len(words)

def split_author_descriptor(text, style_id):
    t=clean_text(text)
    # Strip trailing comma-only after author name
    if t.endswith(','):
        base=clean_author_name(t[:-1])
        if looks_like_person_name(base):
            return [(base, 'AUTOR')]
        return [(base, 'pip')] if base else []
    if ',' in t and (style_id=='AUTOR' or looks_like_person_name(t.split(',',1)[0])):
        first,rest=t.split(',',1)
        first=clean_author_name(first)
        rest=clean_text(rest)
        items=[]
        if first:
            items.append((first,'AUTOR'))
        if rest:
            items.append((rest,'pip'))
        return items
    if style_id=='AUTOR' or looks_like_person_name(t):
        return [(clean_author_name(t), 'AUTOR')]
    return [(t, 'pip')]

def copy_ppr_from_style(p, style_id):
    # applying style is enough; remove direct paragraph indent, alignment overridden as styles dictate
    try: p.style = style_id
    except Exception: pass


def make_paragraph(text='', style_id='a0'):
    p = OxmlElement('w:p')
    if style_id:
        pPr=OxmlElement('w:pPr')
        pStyle=OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    if text:
        r=OxmlElement('w:r')
        t=OxmlElement('w:t')
        if text.startswith(' ') or text.endswith(' '):
            t.set(qn('xml:space'),'preserve')
        t.text=text
        r.append(t)
        p.append(r)
    return p

def replace_paragraph_range(doc, start_idx, end_idx, items):
    # items list of (text, style_id). Replaces paragraphs start_idx..end_idx inclusive.
    paras=doc.paragraphs
    start_el=paras[start_idx]._p
    parent=start_el.getparent()
    insert_pos=parent.index(start_el)
    for idx in range(start_idx, end_idx+1):
        el=paras[idx]._p
        if el.getparent() is not None:
            el.getparent().remove(el)
    for offset,(text,style_id) in enumerate(items):
        parent.insert(insert_pos+offset, make_paragraph(text, style_id))

def find_frontmatter_ranges(doc):
    paras=doc.paragraphs
    ranges=[]
    i=0
    while i < len(paras):
        p=paras[i]
        t=p.text.strip()
        sid=p.style.style_id
        if sid=='UDC' or re.match(r'^(УДК|UDC|DOI)\b', t, flags=re.I):
            # Skip TOC maybe? It is table, not paragraphs. Good.
            start=i
            # Find annotation/abstract before next reference/body. Limit 80 paragraphs
            ann=None
            j=i+1
            while j < len(paras) and j < i+80:
                tj=paras[j].text.strip()
                sj=paras[j].style.style_id
                if sj=='SECTION' or (j>i+2 and (sj=='UDC' or re.match(r'^(УДК|UDC|DOI)\b', tj, flags=re.I))):
                    break
                if re.match(r'^(Анотац(ія|iя)|Abstract|ABSTRACTS?|Annotation)\s*[:.]', tj, flags=re.I):
                    ann=j
                    break
                j+=1
            if ann:
                ranges.append((start, ann-1))
                i=ann+1
                continue
        i+=1
    return ranges

def rebuild_frontmatter(doc):
    report=[]
    # Process from bottom to top so indices stay valid
    ranges=find_frontmatter_ranges(doc)
    for start,end in reversed(ranges):
        paras=doc.paragraphs
        block=paras[start:end+1]
        udc=[]; titles=[]; headers=[]; removed=[]
        for p in block:
            raw=p.text
            text=clean_text(raw)
            sid=p.style.style_id
            if not text:
                continue
            if is_section_note(text):
                removed.append({'type':'section_note','text':text})
                continue
            if is_personal_email_line(text):
                removed.append({'type':'email','text':text})
                continue
            if is_phone_line(text):
                removed.append({'type':'phone','text':text})
                continue
            if sid=='UDC' or re.match(r'^(УДК|UDC|DOI)\b', text, flags=re.I):
                udc.append(text)
            elif sid=='11':
                titles.append(text)
            else:
                headers.append((text,sid))
        # Title might be split over multiple paragraphs: preserve title paragraphs but in canonical order.
        # Header cleanup/splitting
        header_items=[]
        for text,sid in headers:
            for t,style in split_author_descriptor(text,sid):
                if t:
                    header_items.append((t,style))
        # If no titles, leave as is? Not expected.
        items=[]
        for u in udc:
            items.append((u,'UDC'))
        items.append(('', 'a0'))
        for t,style in header_items:
            items.append((t, style))
        if header_items:
            items.append(('', 'a0'))
        for t in titles:
            items.append((t,'11'))
        items.append(('', 'a0'))
        replace_paragraph_range(doc, start, end, items)
        report.append({'start':start,'end':end,'udc':udc,'titles':titles,'headers':[x[0] for x in header_items],'removed':removed})
    return list(reversed(report))

# Table of contents rebuild will happen after render+page extraction.
def get_doc_paragraphs(doc):
    return [(i,p.text.strip(),p.style.style_id) for i,p in enumerate(doc.paragraphs)]

if __name__=='__main__':
    doc=Document(IN)
    report=rebuild_frontmatter(doc)
    # Remove author-supplied section notes anywhere in body/frontmatter (not real SECTION headings).
    removed_notes=[]
    for p in list(doc.paragraphs):
        if p.style.style_id!='SECTION' and is_section_note(clean_text(p.text)):
            removed_notes.append(clean_text(p.text))
            p._element.getparent().remove(p._element)
    doc.save(OUT)
    # Check email removal in doc XML (exclude journal@naukainfo.com)
    with zipfile.ZipFile(OUT) as z:
        xml=z.read('word/document.xml').decode('utf-8')
    emails=[]
    for m in EMAIL_RE.finditer(xml):
        e=(m.group(1)+'@'+m.group(2)).lower()
        if e not in PERSONAL_EMAIL_EXCEPTIONS:
            emails.append(e)
    qa={'input':str(IN),'output':str(OUT),'frontmatter_blocks':len(report),'changes':report,'removed_section_notes_global':removed_notes,'remaining_personal_emails':emails}
    QA.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding='utf-8')
    print(json.dumps({'out':str(OUT),'qa':str(QA),'blocks':len(report),'remaining_personal_emails':emails}, ensure_ascii=False, indent=2))
