from __future__ import annotations
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil, tempfile
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A='http://schemas.openxmlformats.org/drawingml/2006/main'
NS={'w':W,'a':A}

def set_para(p):
    pf=p.paragraph_format
    pf.line_spacing=1.0
    pf.line_spacing_rule=WD_LINE_SPACING.SINGLE
    pf.space_before=Pt(0)
    pf.space_after=Pt(0)
    for run in p.runs:
        run.font.size=Pt(11)

def process_table(tbl):
    seen=set()
    for row in tbl.rows:
        for cell in row.cells:
            # merged cells can repeat; harmless, but avoid duplicate work
            key=id(cell._tc)
            if key in seen: continue
            seen.add(key)
            for p in cell.paragraphs: set_para(p)
            for t in cell.tables: process_table(t)

def normalize_standard(src:Path,tmp:Path):
    doc=Document(src)
    for p in doc.paragraphs: set_para(p)
    for tbl in doc.tables: process_table(tbl)
    for sec in doc.sections:
        for part in [sec.header, sec.footer, sec.first_page_header, sec.first_page_footer,
                     sec.even_page_header, sec.even_page_footer]:
            for p in part.paragraphs: set_para(p)
            for tbl in part.tables: process_table(tbl)
    # Do not rewrite every style. Direct formatting is deterministic and avoids
    # changing shape/table behavior through style inheritance.
    doc.save(tmp)

def ensure_w_size(rpr):
    for tag in [qn('w:sz'), qn('w:szCs')]:
        el=rpr.find(tag)
        if el is None:
            el=OxmlElement(tag.split('}')[-1] if False else ('w:sz' if tag==qn('w:sz') else 'w:szCs'))
            rpr.append(el)
        el.set(qn('w:val'),'22')

def patch_xml(name: str, data: bytes):
    if not name.endswith('.xml'): return data,0
    try: root=etree.fromstring(data)
    except Exception: return data,0
    changed=0
    # Text inside Word text boxes/shapes, including VML fallback copies.
    for tx in root.xpath('.//w:txbxContent',namespaces=NS):
        for r in tx.xpath('.//w:r',namespaces=NS):
            rpr=r.find('{%s}rPr'%W)
            if rpr is None:
                rpr=etree.Element('{%s}rPr'%W)
                r.insert(0,rpr)
            for local in ('sz','szCs'):
                el=rpr.find('{%s}%s'%(W,local))
                if el is None:
                    el=etree.SubElement(rpr,'{%s}%s'%(W,local))
                if el.get('{%s}val'%W)!='22':
                    el.set('{%s}val'%W,'22'); changed+=1
    # SmartArt / DrawingML text: 11 pt is stored as 1100 (1/100 pt).
    if name.startswith('word/diagrams/') or name.startswith('word/drawings/'):
        for el in root.xpath('.//a:rPr | .//a:endParaRPr | .//a:defRPr',namespaces=NS):
            if el.get('sz')!='1100':
                el.set('sz','1100'); changed+=1
    if changed:
        return etree.tostring(root,xml_declaration=True,encoding='UTF-8',standalone=None),changed
    return data,0

def patch_package(tmp:Path,out:Path):
    total=0
    with ZipFile(tmp,'r') as zin, ZipFile(out,'w',ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data=zin.read(info.filename)
            data,n=patch_xml(info.filename,data)
            total+=n
            zout.writestr(info,data)
    return total

def main(src,out):
    src=Path(src); out=Path(out)
    tmp=out.with_suffix('.standard.tmp.docx')
    normalize_standard(src,tmp)
    n=patch_package(tmp,out)
    tmp.unlink(missing_ok=True)
    print({'output':str(out),'shape_text_size_patches':n})

if __name__=='__main__':
    import argparse
    ap=argparse.ArgumentParser()
    ap.add_argument('src'); ap.add_argument('out')
    a=ap.parse_args(); main(a.src,a.out)
