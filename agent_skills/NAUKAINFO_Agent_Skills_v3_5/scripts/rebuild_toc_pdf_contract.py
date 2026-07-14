"""Rebuild NAUKAINFO TABLE OF CONTENTS to match the PDF/Word 2010 contract.

Contract summary:
- Real 3-column borderless Word table under TABLE OF CONTENTS.
- Section = merged row across all 3 columns, styleId TabSEC/visible style Tab_SEC.
- Each article = two rows:
  1) number cell (`1.`), author cell (TabPIP/Tab_PIP), page cell right aligned (TabTaitl/Tab_Taitl)
  2) blank number cell, title cell (TabTaitl/Tab_Taitl), blank page cell
- Do not place the article title in the same row as the author.
- Do not repeat section text in every cell.
- Page numbers must be materialized only after render-verified pagination.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

TITLE_STYLE_ID = "11"  # visible style name: Назва1 in the NAUKAINFO template
DEFAULT_GRID_TWIPS = [600, 8300, 739]


def style_id(p: Paragraph) -> str | None:
    return getattr(p.style, "style_id", None)


def norm(text: str) -> str:
    return " ".join((text or "").split())


def iter_blocks(doc: Document) -> Iterable[tuple[str, Paragraph | Table]]:
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "tbl", Table(child, doc)


def find_toc_and_body_start(doc: Document) -> tuple[int, int, list[tuple[str, Paragraph | Table]]]:
    blocks = list(iter_blocks(doc))
    toc_idx = next(i for i, (k, o) in enumerate(blocks) if k == "p" and norm(o.text).upper() == "TABLE OF CONTENTS")
    body_start = next(i for i, (k, o) in enumerate(blocks[toc_idx + 1 :], toc_idx + 1) if k == "p" and style_id(o) == "SECTION")
    return toc_idx, body_start, blocks


def collect_records(doc: Document) -> list[dict]:
    toc_idx, body_start, blocks = find_toc_and_body_start(doc)
    old_pages: dict[str, str] = {}
    for k, o in blocks[toc_idx + 1 : body_start]:
        if k == "tbl":
            for row in o.rows:
                texts = [norm(c.text) for c in row.cells]
                if len(texts) >= 3 and texts[-1].isdigit():
                    # v2.2 fallback: row could be [number+author, title, page]
                    old_pages[texts[1]] = texts[-1]
    records: list[dict] = []
    current_section = None
    cur = None
    for k, o in blocks[body_start:]:
        if k != "p":
            continue
        sid = style_id(o)
        text = norm(o.text)
        if not text:
            continue
        if sid == "SECTION":
            current_section = text
            continue
        if sid == "AUTOR":
            if cur and cur.get("title"):
                records.append(cur)
                cur = None
            if cur and not cur.get("title"):
                cur["authors"].append(text)
            else:
                cur = {"section": current_section, "authors": [text], "title": None, "page": None}
            continue
        if sid == TITLE_STYLE_ID and cur:
            cur["title"] = text
            cur["page"] = old_pages.get(text)
            continue
        if sid == "REF-TITLE" and cur and cur.get("title"):
            records.append(cur)
            cur = None
    if cur and cur.get("title"):
        records.append(cur)
    # De-duplicate exact records without changing order.
    seen = set()
    clean = []
    for r in records:
        key = (r["section"], tuple(r["authors"]), r["title"])
        if key not in seen:
            seen.add(key)
            clean.append(r)
    return clean


def remove_old_toc_body(doc: Document) -> Paragraph:
    toc_idx, body_start, blocks = find_toc_and_body_start(doc)
    toc_para = blocks[toc_idx][1]
    for k, o in blocks[toc_idx + 1 : body_start]:
        o._element.getparent().remove(o._element)
    return toc_para


def set_table_borders_nil(table: Table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "nil")
    for tag, attrs in [
        ("w:tblW", {"w": str(sum(DEFAULT_GRID_TWIPS)), "type": "dxa"}),
        ("w:tblLayout", {"type": "fixed"}),
        ("w:jc", {"val": "center"}),
    ]:
        el = tblPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tblPr.insert(0, el)
        for k, v in attrs.items():
            el.set(qn("w:" + k), v)


def set_cell_width(cell, width: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width))
    tcW.set(qn("w:type"), "dxa")


def set_grid(table: Table, widths: list[int] = DEFAULT_GRID_TWIPS) -> None:
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells[: len(widths)]):
            set_cell_width(cell, widths[idx])


def set_cell_text(cell, text: str, style_name: str, align=None) -> Paragraph:
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.style = style_name
    if align is not None:
        p.alignment = align
    p.add_run(text)
    return p


def build_toc_table(doc: Document, records: list[dict]) -> Table:
    rows = 0
    last_section = None
    for r in records:
        if r["section"] != last_section:
            rows += 1
            last_section = r["section"]
        rows += 2
    table = doc.add_table(rows=rows, cols=3)
    table.autofit = False
    set_table_borders_nil(table)
    set_grid(table)
    row_idx = 0
    last_section = None
    for number, r in enumerate(records, 1):
        if r["section"] != last_section:
            merged = table.rows[row_idx].cells[0].merge(table.rows[row_idx].cells[1]).merge(table.rows[row_idx].cells[2])
            set_cell_text(merged, r["section"], "Tab_SEC", WD_ALIGN_PARAGRAPH.CENTER)
            last_section = r["section"]
            row_idx += 1
        set_cell_text(table.rows[row_idx].cells[0], f"{number}.", "Tab_Taitl", WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[row_idx].cells[1], ", ".join(r["authors"]), "Tab_PIP", WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[row_idx].cells[2], str(r.get("page") or "?"), "Tab_Taitl", WD_ALIGN_PARAGRAPH.RIGHT)
        row_idx += 1
        set_cell_text(table.rows[row_idx].cells[0], "", "Tab_Taitl", WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[row_idx].cells[1], r["title"], "Tab_Taitl", WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[row_idx].cells[2], "", "Tab_Taitl", WD_ALIGN_PARAGRAPH.RIGHT)
        row_idx += 1
    set_grid(table)
    return table


def rebuild_docx_toc(input_docx: str | Path, output_docx: str | Path, qa_json: str | Path | None = None) -> list[dict]:
    doc = Document(str(input_docx))
    records = collect_records(doc)
    toc_para = remove_old_toc_body(doc)
    table = build_toc_table(doc, records)
    toc_para._p.addnext(table._tbl)
    doc.save(str(output_docx))
    if qa_json:
        Path(qa_json).write_text(json.dumps({
            "records": records,
            "row_contract": "section merged row; article author row; article title row",
            "style_ids": {"section":"TabSEC", "author":"TabPIP", "title":"TabTaitl", "number":"TabTaitl", "page":"TabTaitl"},
            "grid_twips": DEFAULT_GRID_TWIPS,
        }, ensure_ascii=False, indent=2), encoding="utf-8")
    return records


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx")
    parser.add_argument("output_docx")
    parser.add_argument("--qa_json")
    args = parser.parse_args()
    rebuild_docx_toc(args.input_docx, args.output_docx, args.qa_json)
