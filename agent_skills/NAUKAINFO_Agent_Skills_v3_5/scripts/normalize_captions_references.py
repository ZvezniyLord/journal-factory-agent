from __future__ import annotations

import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt

TAIL_MARKER = 'SCIENCE \nIN THE MODERN WORLD'
REF_STAMP_UA = 'СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ:'


def iter_block_items(parent):
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def has_drawing(p: Paragraph) -> bool:
    return bool(p._p.xpath('.//w:drawing | .//w:pict | .//w:object'))


def remove_paragraph(p: Paragraph) -> None:
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None


def insert_paragraph_before(paragraph: Paragraph, text: str = '') -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def clear_direct_child(pPr, tag: str) -> None:
    if pPr is None:
        return
    el = pPr.find(qn(tag))
    if el is not None:
        pPr.remove(el)


def set_run_11(run, bold=None, italic=None):
    run.font.size = Pt(11)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_common(p: Paragraph, *, align=None, first_line_zero=True, before=0, after=0, keep_next=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if first_line_zero:
        pf.first_line_indent = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if keep_next is not None:
        pf.keep_with_next = keep_next


def apply_style_exact(p: Paragraph, style: str, *, align=None, first_line_zero=True, keep_next=None):
    p.style = style
    set_para_common(p, align=align, first_line_zero=first_line_zero, keep_next=keep_next)
    for r in p.runs:
        set_run_11(r)


def looks_like_person_name(text: str) -> bool:
    t = re.sub(r'\s+', ' ', text.strip())
    if not t:
        return False
    low = t.lower()
    blockers = ['університет','інститут','академ','кафедр','здобувач','керівник','доцент','професор','hon.','phd','доктор','кандидат','м. ','україна','orcid','спілк','правління','діяч']
    if any(x in low for x in blockers):
        return False
    words = t.split()
    if len(words) not in (2,3,4):
        return False
    return all(re.match(r"^[А-ЯІЇЄҐA-Z][А-Яа-яІіЇїЄєҐґA-Za-z'’\-]+$", w) for w in words)


def is_article_title(text: str) -> bool:
    t = text.strip()
    letters = [c for c in t if c.isalpha()]
    if len(letters) < 25:
        return False
    return sum(c.isupper() for c in letters) / len(letters) > 0.86


def next_num_id(numbering_root) -> int:
    vals=[]
    for el in numbering_root.findall(qn('w:num')):
        v=el.get(qn('w:numId'))
        if v and v.isdigit(): vals.append(int(v))
    return max(vals or [0])+1


def create_reference_num(doc: _Document) -> int:
    """Fresh decimal numbering instance from ETALON abstractNumId=1, restarted at 1."""
    root = doc.part.numbering_part.element
    num_id = next_num_id(root)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract = OxmlElement('w:abstractNumId')
    abstract.set(qn('w:val'), '1')
    num.append(abstract)
    override = OxmlElement('w:lvlOverride')
    override.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:startOverride')
    start.set(qn('w:val'), '1')
    override.append(start)
    num.append(override)
    root.append(num)
    return num_id


def set_reference_num(p: Paragraph, num_id: int):
    pPr = p._p.get_or_add_pPr()
    clear_direct_child(pPr, 'w:numPr')
    clear_direct_child(pPr, 'w:ind')
    clear_direct_child(pPr, 'w:tabs')
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'),'0')
    num = OxmlElement('w:numId'); num.set(qn('w:val'),str(num_id))
    numPr.append(ilvl); numPr.append(num)
    pPr.append(numPr)


def normalize_reference_block(doc: _Document, refs_idx: int, tail_idx: int, audit: dict):
    paras = doc.paragraphs
    heading = paras[refs_idx]
    old_stamp = heading.text.strip()
    heading.text = REF_STAMP_UA
    apply_style_exact(heading, 'REF-TITLE', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_zero=True, keep_next=True)
    for r in heading.runs:
        r.bold = None

    # Exactly one blank paragraph before and after the stamp.
    paras = doc.paragraphs
    refs_idx = next(i for i,p in enumerate(paras) if p._p is heading._p)
    before_blanks=[]
    k=refs_idx-1
    while k>=0 and not paras[k].text.strip():
        before_blanks.append(paras[k]); k-=1
    if not before_blanks:
        blank=insert_paragraph_before(heading,'')
        blank.style='Normal'; set_para_common(blank, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_zero=True)
    else:
        keep=before_blanks[0]
        keep.style='Normal'; set_para_common(keep, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_zero=True)
        for extra in before_blanks[1:]: remove_paragraph(extra)

    paras=doc.paragraphs; refs_idx=next(i for i,p in enumerate(paras) if p._p is heading._p)
    after_blanks=[]
    k=refs_idx+1
    while k < len(paras) and not paras[k].text.strip():
        after_blanks.append(paras[k]); k+=1
    if not after_blanks:
        # Insert one blank immediately after heading.
        new_p = OxmlElement('w:p'); heading._p.addnext(new_p)
        blank = Paragraph(new_p, heading._parent); blank.style='Normal'
    else:
        keep=after_blanks[0]; keep.style='Normal'
        for extra in after_blanks[1:]: remove_paragraph(extra)

    # Refresh tail after deletions.
    paras=doc.paragraphs
    tail_idx = next(i for i,p in enumerate(paras[refs_idx+1:], refs_idx+1) if p.text.strip()==TAIL_MARKER)
    entries=[p for p in paras[refs_idx+1:tail_idx] if p.text.strip()]
    num_id=create_reference_num(doc)
    for p in entries:
        p.style='REFER'
        set_reference_num(p,num_id)
        pf=p.paragraph_format
        pf.line_spacing=1.0; pf.line_spacing_rule=WD_LINE_SPACING.SINGLE
        pf.space_before=Pt(0); pf.space_after=Pt(0)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        # Do not set firstLine/left directly: numbering definition provides exact 1 cm hanging geometry.
        for r in p.runs:
            set_run_11(r)
    audit['references']={'old_stamp':old_stamp,'new_stamp':REF_STAMP_UA,'entries':len(entries),'num_id':num_id}


def normalize_table_figure_captions(doc: _Document, article_start_idx: int, refs_idx: int, audit: dict):
    paras=doc.paragraphs
    changes=[]
    for i in range(article_start_idx, refs_idx):
        p=paras[i]; t=p.text.strip()
        if has_drawing(p):
            apply_style_exact(p,'РИС',align=WD_ALIGN_PARAGRAPH.CENTER,first_line_zero=True,keep_next=True)
            changes.append((i,'figure_anchor',t))
        elif re.match(r'^(Рис\.?\s*\d+|Рисунок\s+\d+|Figure\s+\d+)',t,re.I):
            apply_style_exact(p,'РисПід',align=WD_ALIGN_PARAGRAPH.CENTER,first_line_zero=True,keep_next=True)
            for r in p.runs: r.bold=None
            changes.append((i,'figure_caption',t))
        elif re.match(r'^(Таблиця|Table)\s*\d+',t,re.I):
            p.style='Normal'; set_para_common(p,align=WD_ALIGN_PARAGRAPH.RIGHT,first_line_zero=True,keep_next=True)
            for r in p.runs: set_run_11(r,bold=True)
            changes.append((i,'table_number',t))
            # Next nonblank paragraph is table title.
            j=i+1
            while j<refs_idx and not paras[j].text.strip(): j+=1
            if j<refs_idx:
                title=paras[j]
                title.style='Normal'; set_para_common(title,align=WD_ALIGN_PARAGRAPH.CENTER,first_line_zero=True,keep_next=True)
                for r in title.runs: set_run_11(r,bold=True)
                changes.append((j,'table_title',title.text.strip()))
        elif re.match(r'^(Джерело|Source)\s*[:.]',t,re.I):
            p.style='Normal'; set_para_common(p,first_line_zero=True,keep_next=False)
            for r in p.runs: set_run_11(r)
            changes.append((i,'source_note',t))

    # Canonical table-cell text, preserving alignment/emphasis.
    table_para_count=0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    old_align=p.alignment
                    p.style='TABLETEXT'
                    set_para_common(p,align=old_align,first_line_zero=True)
                    for r in p.runs: set_run_11(r)
                    table_para_count+=1
    audit['captions']={'changes':changes,'table_paragraphs':table_para_count}


def apply_article_styles(doc: _Document, section: str, audit: dict):
    paras=doc.paragraphs
    toc_idx=next(i for i,p in enumerate(paras) if p.text.strip()=='TABLE OF CONTENTS')
    tail_idx=next(i for i,p in enumerate(paras[toc_idx+1:],toc_idx+1) if p.text.strip()==TAIL_MARKER)
    udc_idx=next(i for i,p in enumerate(paras[toc_idx+1:tail_idx],toc_idx+1) if re.match(r'^(УДК|UDC)\b',p.text.strip(),re.I))
    # DOI may precede UDC; section goes before first service metadata.
    first_meta=udc_idx
    for i in range(udc_idx-1,toc_idx,-1):
        if re.match(r'^(DOI\s*:|https?://doi\.org/)',paras[i].text.strip(),re.I): first_meta=i
        elif paras[i].text.strip(): break
    preceding=next((p for p in reversed(paras[toc_idx+1:first_meta]) if p.text.strip()),None)
    if preceding is None or preceding.text.strip()!=section:
        sec=insert_paragraph_before(paras[first_meta],section)
    else: sec=preceding
    apply_style_exact(sec,'SECTION',align=WD_ALIGN_PARAGRAPH.CENTER,first_line_zero=True,keep_next=True)
    sec.paragraph_format.space_after=Pt(18)
    for r in sec.runs: r.bold=None

    paras=doc.paragraphs
    toc_idx=next(i for i,p in enumerate(paras) if p.text.strip()=='TABLE OF CONTENTS')
    tail_idx=next(i for i,p in enumerate(paras[toc_idx+1:],toc_idx+1) if p.text.strip()==TAIL_MARKER)
    sec_idx=next(i for i,p in enumerate(paras[toc_idx+1:tail_idx],toc_idx+1) if p.text.strip()==section)
    udc_idx=next(i for i,p in enumerate(paras[sec_idx+1:tail_idx],sec_idx+1) if re.match(r'^(УДК|UDC)\b',p.text.strip(),re.I))
    title_idx=next(i for i,p in enumerate(paras[udc_idx+1:tail_idx],udc_idx+1) if is_article_title(p.text))
    refs_idx=next(i for i,p in enumerate(paras[title_idx+1:tail_idx],title_idx+1) if p.text.strip().upper().startswith('СПИСОК ВИКОРИСТАН'))

    for i in range(sec_idx+1,title_idx):
        p=paras[i]; t=p.text.strip()
        if not t: continue
        if re.match(r'^(DOI\s*:|https?://doi\.org/|УДК\b|UDC\b)',t,re.I):
            apply_style_exact(p,'UDC',align=WD_ALIGN_PARAGRAPH.LEFT,first_line_zero=True)
            for r in p.runs: r.bold=None
        elif looks_like_person_name(t):
            apply_style_exact(p,'AUTOR',align=WD_ALIGN_PARAGRAPH.RIGHT,first_line_zero=True)
        else:
            apply_style_exact(p,'pip',align=WD_ALIGN_PARAGRAPH.RIGHT,first_line_zero=True)

    apply_style_exact(paras[title_idx],'Назва1',align=WD_ALIGN_PARAGRAPH.CENTER,first_line_zero=True,keep_next=True)
    # Annotation/keywords retain Normal body indent; list items cannot inherit a positive body first-line indent.
    for i in range(title_idx+1,refs_idx):
        p=paras[i]; t=p.text.strip()
        if p._p.pPr is not None and p._p.pPr.numPr is not None and not t.upper().startswith('СПИСОК'):
            # ordinary in-body lists keep numbering but not a positive first-line indent
            pf=p.paragraph_format
            if pf.first_line_indent is not None and pf.first_line_indent.pt>0: pf.first_line_indent=Pt(0)
            pf.line_spacing=1.0; pf.space_before=Pt(0); pf.space_after=Pt(0)

    audit['article']={'section':section,'sec_idx':sec_idx,'udc_idx':udc_idx,'title_idx':title_idx,'refs_idx':refs_idx,'tail_idx':tail_idx}
    return sec_idx,title_idx,refs_idx,tail_idx


def audit_output(path: Path, section: str) -> dict:
    d=Document(path); paras=d.paragraphs
    out={'file':str(path),'paragraphs':len(paras),'tables':len(d.tables),'section_present':False,'stamps':[],'refs':[],'captions':[],'issues':[]}
    for i,p in enumerate(paras):
        t=p.text.strip()
        if t==section: out['section_present']=p.style.name=='SECTION'
        if t.upper().startswith('СПИСОК ВИКОРИСТАН'):
            out['stamps'].append({'i':i,'text':t,'style':p.style.name,'before_blank': i>0 and not paras[i-1].text.strip(),'after_blank': i+1<len(paras) and not paras[i+1].text.strip()})
        if p.style.name=='REFER':
            pPr=p._p.pPr; num=None; ind=None
            if pPr is not None and pPr.numPr is not None and pPr.numPr.numId is not None: num=pPr.numPr.numId.val
            if pPr is not None and pPr.ind is not None: ind=pPr.ind.xml
            out['refs'].append({'i':i,'numId':num,'direct_ind':ind,'align':str(p.alignment),'text':t[:90]})
        if re.match(r'^(Рис\.|Рисунок|Figure|Таблиця|Table|Джерело|Source)',t,re.I):
            pf=p.paragraph_format
            out['captions'].append({'i':i,'text':t[:100],'style':p.style.name,'align':str(p.alignment),'first':None if pf.first_line_indent is None else pf.first_line_indent.twips,'keep_next':pf.keep_with_next})
    if not out['section_present']: out['issues'].append('section style missing')
    for s in out['stamps']:
        if s['text']!=REF_STAMP_UA: out['issues'].append('wrong reference stamp')
        if not s['before_blank']: out['issues'].append('missing blank before reference stamp')
        if not s['after_blank']: out['issues'].append('missing blank after reference stamp')
        if s['style']!='REF-TITLE': out['issues'].append('reference heading style')
    for r in out['refs']:
        if r['direct_ind'] is not None: out['issues'].append('direct indent on reference')
        if r['numId'] is None: out['issues'].append('reference numbering missing')
    return out


def process(src: Path, out: Path, section: str):
    doc=Document(src)
    audit={}
    sec_idx,title_idx,refs_idx,tail_idx=apply_article_styles(doc,section,audit)
    # Recompute after possible section insertion.
    paras=doc.paragraphs
    refs_idx=next(i for i,p in enumerate(paras) if p.text.strip().upper().startswith('СПИСОК ВИКОРИСТАН'))
    tail_idx=next(i for i,p in enumerate(paras[refs_idx+1:],refs_idx+1) if p.text.strip()==TAIL_MARKER)
    normalize_table_figure_captions(doc,title_idx,refs_idx,audit)
    normalize_reference_block(doc,refs_idx,tail_idx,audit)
    doc.save(out)
    qa=audit_output(out,section)
    audit['qa']=qa
    out.with_suffix('.audit.json').write_text(json.dumps(audit,ensure_ascii=False,indent=2),encoding='utf-8')
    if qa['issues']:
        raise RuntimeError('QA failed: '+repr(qa['issues']))
    print(json.dumps(audit,ensure_ascii=False,indent=2))


if __name__=='__main__':
    import argparse
    ap=argparse.ArgumentParser()
    ap.add_argument('src'); ap.add_argument('out'); ap.add_argument('--section',required=True)
    a=ap.parse_args(); process(Path(a.src),Path(a.out),a.section)
