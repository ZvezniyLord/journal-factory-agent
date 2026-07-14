#!/usr/bin/env python3
"""Create a compact, article-scoped UDC lookup request.

No code is invented offline. In a multi-article journal, a UDC from another
article must never satisfy the target article. Select the target by title or
article number.
"""
from __future__ import annotations
import argparse, json, re
from pathlib import Path
from docx import Document
from udc_article_gate import article_records, header_bounds, is_title, is_udc_marker

ANNOT_RE = re.compile(r"^(?:Анотація|Abstract|Annotation)\s*[:.]", re.I)
KEY_RE = re.compile(r"^(?:Ключові слова|Keywords?)\s*[:.]", re.I)


def _select_article(doc, article_title: str | None, article_number: int | None):
    paragraphs = list(doc.paragraphs)
    records = article_records(paragraphs)
    if records:
        if article_title:
            needle = " ".join(article_title.split()).casefold()
            hits = [
                r for r in records
                if needle in " ".join(r["title"].split()).casefold()
                or " ".join(r["title"].split()).casefold() in needle
            ]
            if len(hits) != 1:
                return None, {
                    "status": "ARTICLE_SELECTOR_AMBIGUOUS",
                    "match_count": len(hits),
                    "article_title": article_title,
                    "available_titles": [r["title"] for r in records],
                }
            return hits[0], None
        if article_number is not None:
            hits = [r for r in records if r["article_no"] == article_number]
            if len(hits) != 1:
                return None, {
                    "status": "ARTICLE_SELECTOR_NOT_FOUND",
                    "article_number": article_number,
                    "available_article_numbers": [r["article_no"] for r in records],
                }
            return hits[0], None
        if len(records) > 1:
            return None, {
                "status": "ARTICLE_SELECTOR_REQUIRED",
                "article_count": len(records),
                "available_titles": [r["title"] for r in records],
            }
        return records[0], None

    # Source article without template styles: treat the whole file as one article.
    return {
        "article_no": 1,
        "title_index": None,
        "title": "",
        "header_start": 0,
        "header_end": len(paragraphs),
        "udc_markers": [p.text.strip() for p in paragraphs if is_udc_marker(p.text)],
    }, None


def build_request(
    docx: Path,
    article_id: str,
    section: str,
    article_title: str | None = None,
    article_number: int | None = None,
) -> dict:
    doc = Document(docx)
    paragraphs = list(doc.paragraphs)
    rec, error = _select_article(doc, article_title, article_number)
    if error:
        error.update({"article_id": article_id, "needs_operator_review": True})
        return error

    existing = rec.get("udc_markers", [])
    if existing:
        return {
            "status": "existing_udc",
            "article_id": article_id,
            "article_no": rec.get("article_no"),
            "title": rec.get("title", ""),
            "udc": existing[0],
            "needs_operator_review": False,
        }

    if rec.get("title_index") is not None:
        title_index = rec["title_index"]
        title = paragraphs[title_index].text.strip()
        next_title = next(
            (i for i in range(title_index + 1, len(paragraphs)) if is_title(paragraphs[i])),
            len(paragraphs),
        )
        body = [p.text.strip() for p in paragraphs[title_index + 1 : next_title] if p.text.strip()]
    else:
        texts = [p.text.strip() for p in paragraphs if p.text.strip()]
        title = next(
            (t for t in texts if len(t) > 20 and t.upper() == t and not t.startswith("СПИСОК")),
            "",
        )
        body = texts

    annotation = next((t for t in body if ANNOT_RE.match(t)), "")
    keywords = next((t for t in body if KEY_RE.match(t)), "")
    excerpt = next(
        (t for t in body if t not in {annotation, keywords} and len(t) >= 120),
        "",
    )[:900]

    return {
        "status": "UDC_LOOKUP_REQUIRED",
        "article_id": article_id,
        "article_no": rec.get("article_no"),
        "section": section,
        "title": title,
        "annotation": annotation,
        "keywords": keywords,
        "body_excerpt": excerpt,
        "instruction": (
            "Run online authoritative UDC research now. Return one primary code, "
            "up to three alternatives, confidence, evidence sources, and a review flag. "
            "A high-confidence documented candidate may be inserted and must be tagged "
            "udc_source=generated; ambiguous classification remains blocked."
        ),
        "needs_operator_review": True,
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    ap.add_argument("--article-id", required=True)
    ap.add_argument("--section", required=True)
    ap.add_argument("--article-title")
    ap.add_argument("--article-number", type=int)
    ap.add_argument("--out", type=Path, required=True)
    a = ap.parse_args()
    result = build_request(
        a.docx,
        a.article_id,
        a.section,
        article_title=a.article_title,
        article_number=a.article_number,
    )
    a.out.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
