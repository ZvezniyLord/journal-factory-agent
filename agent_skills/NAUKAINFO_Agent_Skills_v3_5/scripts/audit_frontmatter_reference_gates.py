#!/usr/bin/env python3
"""Fail-closed audit for NAUKAINFO frontmatter and bibliography business gates."""
from __future__ import annotations
import argparse, json, re
from pathlib import Path
from docx import Document
from udc_article_gate import audit_udc_per_article

TERMINAL = re.compile(r"[,;:]\s*$")
MANUAL_NUM = re.compile(r"^\s*(?:\(?\d+\)?[.)]?|\d+)\s+")
FRONT_STYLES = {"UDC", "AUTOR", "pip", "Назва1"}

def has_numpr(p):
    return p._p.pPr is not None and p._p.pPr.numPr is not None

def has_forbidden_geometry(p):
    ppr=p._p.pPr
    if ppr is None: return False
    return any(getattr(ppr, name, None) is not None for name in ("numPr","ind","tabs","outlineLvl"))

def audit(path: Path):
    d=Document(path)
    defects=[]
    for i,p in enumerate(d.paragraphs):
        sid=p.style.style_id
        text=p.text.strip()
        if sid=="AUTOR" and text and "\n" not in text and TERMINAL.search(text):
            defects.append({"code":"AUTHOR_TERMINAL_PUNCTUATION","paragraph":i,"text":text})
        if sid in FRONT_STYLES and text and has_forbidden_geometry(p):
            defects.append({"code":"FRONTMATTER_LIST_OR_INDENT","paragraph":i,"style":sid,"text":text})
        if sid=="REFER" and has_numpr(p) and MANUAL_NUM.match(text):
            defects.append({"code":"REFERENCE_DUPLICATE_NUMBER","paragraph":i,"text":text})
    # Mandatory UDC is audited per article title. A UDC elsewhere in the journal,
    # or a DOI paragraph that happens to use the UDC style, cannot satisfy it.
    defects.extend(audit_udc_per_article(d.paragraphs))
    # TOC author cells: numeric first column marks article rows.
    if d.tables:
        for ri,row in enumerate(d.tables[0].rows):
            if re.fullmatch(r"\d+\.", row.cells[0].text.strip()):
                text=" ".join(row.cells[1].text.split())
                for name in [x.strip() for x in text.split(';') if x.strip()]:
                    if TERMINAL.search(name):
                        defects.append({"code":"TOC_AUTHOR_TERMINAL_PUNCTUATION","row":ri,"text":name})
    return {"file":str(path),"status":"PASS" if not defects else "FAIL","defect_count":len(defects),"defects":defects}

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument('docx',type=Path)
    ap.add_argument('--out-json',type=Path)
    a=ap.parse_args()
    report=audit(a.docx)
    if a.out_json: a.out_json.write_text(json.dumps(report,ensure_ascii=False,indent=2),encoding='utf-8')
    print(json.dumps(report,ensure_ascii=False,indent=2))
    raise SystemExit(0 if report['status']=='PASS' else 2)
if __name__=='__main__': main()
