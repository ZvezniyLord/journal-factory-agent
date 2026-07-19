from __future__ import annotations

from pathlib import Path
from typing import Any
import locale
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.section import CT_SectPr
from docxcompose.composer import Composer


TOC_STYLE_MAPPING = (("SECTION", 1), ("Назва1", 2))


def _toc_list_separator() -> str:
    try:
        import winreg

        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Control Panel\International") as key:
            separator = str(winreg.QueryValueEx(key, "sList")[0]).strip()
            if separator:
                return separator
    except (ImportError, OSError):
        pass
    return ";" if locale.localeconv().get("decimal_point") == "," else ","


def toc_instruction() -> str:
    separator = _toc_list_separator()
    mapping = separator.join(
        value for style_name, level in TOC_STYLE_MAPPING for value in (style_name, str(level))
    )
    return f'TOC \\h \\z \\t "{mapping}"'


def compose_articles_into_etalon(
    master_path: Path,
    articles: list[dict[str, Any]],
    output_path: Path,
    pagination_policy: dict[str, Any] | None = None,
    layout_adjustments: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Insert ordered article DOCX files after TOC and before the tail section."""
    master = Document(str(master_path))
    expected_section_count = len(master.sections)
    _ensure_toc_field(master)
    composer = Composer(master, preserve_styles=False)
    insert_index = find_article_insert_index(master)
    initial_index = insert_index
    inserted: list[dict[str, Any]] = []
    blockers: list[str] = []
    previous_section: str | None = None

    for position, article in enumerate(articles, start=1):
        source = Path(article["source_file"])
        if source.suffix.lower() != ".docx":
            blockers.append(f"unsupported_article_extension:{source}")
            continue
        if not source.is_file():
            blockers.append(f"article_source_missing:{source}")
            continue

        section = str(article.get("section") or "").strip()
        separator = _separator_document(
            add_page_break=True,
            section_title=section if section and section != previous_section else None,
            block_bookmark=f"JF_LAYOUT_ARTICLE_{position:03d}_START",
        )
        separator_count = _inserted_element_count(separator)
        composer.insert(insert_index, separator)
        insert_index += separator_count

        source_doc = Document(str(source))
        article_count = _inserted_element_count(source_doc)
        composer.insert(insert_index, source_doc)
        insert_index += article_count

        inserted.append(
            {
                "position": position,
                "article_id": article.get("article_id"),
                "source_file": str(source),
                "section": section,
                "title_bookmark": article.get("title_bookmark"),
                "inserted_body_elements": article_count,
            }
        )
        if section:
            previous_section = section

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if blockers:
        return {
            "status": "BLOCKED",
            "output": None,
            "toc_tail_insert_index": initial_index,
            "inserted": inserted,
            "blockers": blockers,
        }

    composer.save(str(output_path))
    structure_report = _finalize_document_structure(
        output_path,
        articles,
        expected_section_count,
        pagination_policy=pagination_policy,
        layout_adjustments=layout_adjustments,
    )
    if structure_report["status"] != "PASS":
        return {
            "status": "BLOCKED",
            "output": str(output_path),
            "toc_tail_insert_index": initial_index,
            "inserted": inserted,
            "inserted_count": len(inserted),
            "structure": structure_report,
            "blockers": structure_report["blockers"],
        }
    return {
        "status": "PASS",
        "output": str(output_path),
        "toc_tail_insert_index": initial_index,
        "inserted": inserted,
        "inserted_count": len(inserted),
        "structure": structure_report,
        "blockers": [],
    }


def find_article_insert_index(document: Document) -> int:
    """Find the section-break paragraph after TABLE OF CONTENTS.

    Article blocks are inserted immediately before this paragraph, preserving
    the final service-page section that follows it.
    """
    toc_seen = False
    for index, element in enumerate(document.element.body):
        if isinstance(element, CT_SectPr):
            continue
        text = "".join(element.xpath(".//w:t/text()"))
        normalized = " ".join(text.upper().split())
        if "TABLE OF CONTENTS" in normalized:
            toc_seen = True
            continue
        if toc_seen and element.xpath("./w:pPr/w:sectPr"):
            return index
    raise ValueError("ETALON insertion anchor after TABLE OF CONTENTS was not found")


def _separator_document(
    add_page_break: bool,
    section_title: str | None,
    block_bookmark: str | None = None,
) -> Document:
    separator = Document()
    paragraph = separator.add_paragraph(section_title or "")
    ppr = paragraph._p.get_or_add_pPr()
    if add_page_break:
        page_break_before = OxmlElement("w:pageBreakBefore")
        ppr.append(page_break_before)
    if section_title:
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), "SECTION")
        ppr.insert(0, pstyle)
    else:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "20")
        spacing.set(qn("w:lineRule"), "exact")
        ppr.append(spacing)
        paragraph_rpr = OxmlElement("w:rPr")
        for name in ("w:sz", "w:szCs"):
            size = OxmlElement(name)
            size.set(qn("w:val"), "2")
            paragraph_rpr.append(size)
        ppr.append(paragraph_rpr)
    if block_bookmark:
        match = re.search(r"(\d+)", block_bookmark)
        if match is None:
            raise ValueError(f"invalid_block_bookmark:{block_bookmark}")
        bookmark_id = str(4000 + int(match.group(1)))
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), bookmark_id)
        bookmark_start.set(qn("w:name"), block_bookmark)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), bookmark_id)
        paragraph._p.insert(1, bookmark_start)
        paragraph._p.append(bookmark_end)
    return separator


def _inserted_element_count(document: Document) -> int:
    return sum(1 for element in document.element.body if not isinstance(element, CT_SectPr))


def _ensure_toc_field(document: Document) -> None:
    instructions = " ".join(document.element.body.xpath(".//w:instrText/text()"))
    if re_search_toc(instructions):
        return
    heading = next(
        (
            element
            for element in document.element.body
            if "TABLE OF CONTENTS" in " ".join("".join(element.xpath(".//w:t/text()")).upper().split())
        ),
        None,
    )
    if heading is None:
        raise ValueError("TABLE OF CONTENTS heading was not found")
    paragraph = OxmlElement("w:p")
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {toc_instruction()} "
    instruction_run.append(instruction)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    placeholder_run = OxmlElement("w:r")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table of contents will be updated by Microsoft Word."
    placeholder_run.append(placeholder)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    for element in (begin_run, instruction_run, separate_run, placeholder_run, end_run):
        paragraph.append(element)
    heading.addnext(paragraph)
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def re_search_toc(value: str) -> bool:
    return bool(re.search(r"\bTOC\b", value, flags=re.IGNORECASE))


def _finalize_document_structure(
    output_path: Path,
    articles: list[dict[str, Any]],
    expected_section_count: int,
    pagination_policy: dict[str, Any] | None = None,
    layout_adjustments: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    document = Document(str(output_path))
    expected_bookmark_count = sum(bool(item.get("title_bookmark")) for item in articles)
    title_style_id = _style_id_by_name(document, "Назва1") if expected_bookmark_count else ""
    blockers: list[str] = []
    applied_bookmarks: list[str] = []
    for article in articles:
        bookmark_name = str(article.get("title_bookmark") or "")
        if not bookmark_name:
            continue
        starts = [
            item
            for item in document.element.body.xpath(".//w:bookmarkStart")
            if item.get(qn("w:name")) == bookmark_name
        ]
        if len(starts) != 1:
            blockers.append(f"article_bookmark_count:{bookmark_name}:{len(starts)}")
            continue
        paragraph = starts[0].getparent()
        if paragraph.tag != qn("w:p"):
            blockers.append(f"article_bookmark_not_in_paragraph:{bookmark_name}")
            continue
        ppr = paragraph.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            paragraph.insert(0, ppr)
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is None:
            pstyle = OxmlElement("w:pStyle")
            ppr.insert(0, pstyle)
        pstyle.set(qn("w:val"), title_style_id)
        applied_bookmarks.append(bookmark_name)

    section_properties = document.element.body.xpath(".//w:sectPr")
    if len(section_properties) != expected_section_count:
        blockers.append(
            f"unexpected_section_count:{len(section_properties)}:expected={expected_section_count}"
        )
    for sect_pr in section_properties:
        for reference in list(sect_pr.xpath("./w:headerReference | ./w:footerReference")):
            sect_pr.remove(reference)
        for pg_num in list(sect_pr.xpath("./w:pgNumType")):
            sect_pr.remove(pg_num)
        for title_page in list(sect_pr.xpath("./w:titlePg")):
            sect_pr.remove(title_page)
    if pagination_policy:
        compatibility_mode = pagination_policy.get("word_compatibility_mode")
        if compatibility_mode is not None:
            _set_word_compatibility_mode(document, int(compatibility_mode))
        numbered_section_index = int(
            pagination_policy.get("numbered_section_index") or 1
        )
        first_printed_page = int(pagination_policy.get("first_printed_page") or 1)
        if not 1 <= numbered_section_index <= len(section_properties):
            blockers.append(
                f"numbered_section_index:{numbered_section_index}:sections={len(section_properties)}"
            )
        else:
            numbered = section_properties[numbered_section_index - 1]
            pg_num = OxmlElement("w:pgNumType")
            pg_num.set(qn("w:start"), str(first_printed_page))
            numbered.insert(0, pg_num)

            footer_distance = document.sections[0].footer_distance
            for index, section in enumerate(document.sections, start=1):
                section.footer_distance = footer_distance
                if index < numbered_section_index:
                    section.footer.is_linked_to_previous = False
                    section.footer.paragraphs[0].clear()
                elif index == numbered_section_index:
                    section.footer.is_linked_to_previous = False
                    paragraph = section.footer.paragraphs[0]
                    paragraph.clear()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_complex_field(paragraph, "PAGE")
                else:
                    section.footer.is_linked_to_previous = True
    elif section_properties:
        first = section_properties[0]
        pg_num = OxmlElement("w:pgNumType")
        pg_num.set(qn("w:start"), "1")
        first.insert(0, pg_num)
        first.append(OxmlElement("w:titlePg"))
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_complex_field(paragraph, "PAGE")
        for section in document.sections[1:]:
            section.footer.is_linked_to_previous = True

    layout_report = _apply_layout_adjustments(document, layout_adjustments or [])
    blockers.extend(layout_report["blockers"])

    document.save(str(output_path))
    section_count = len(Document(str(output_path)).sections)
    if len(applied_bookmarks) != expected_bookmark_count:
        blockers.append(
            f"title_style_application_count:{len(applied_bookmarks)}:{expected_bookmark_count}"
        )
    return {
        "status": "PASS" if not blockers else "BLOCKED",
        "toc_instruction": toc_instruction(),
        "toc_style_mapping": dict(TOC_STYLE_MAPPING),
        "toc_style_separator": _toc_list_separator(),
        "title_style_id": title_style_id,
        "title_styles_applied": len(applied_bookmarks),
        "article_bookmarks": applied_bookmarks,
        "section_count": section_count,
        "page_number_policy": pagination_policy or {
            "first_section_start": 1,
            "title_page_number_hidden": True,
            "later_sections_continue": True,
            "footer_field": "PAGE",
        },
        "layout_adjustments": layout_report,
        "blockers": blockers,
    }


def _apply_layout_adjustments(
    document: Document,
    adjustments: list[dict[str, Any]],
) -> dict[str, Any]:
    blockers: list[str] = []
    applied: list[dict[str, Any]] = []
    seen_ordinals: set[int] = set()
    for adjustment in adjustments:
        ordinal = int(adjustment.get("ordinal") or 0)
        if ordinal <= 0 or ordinal in seen_ordinals:
            blockers.append(f"invalid_layout_adjustment_ordinal:{ordinal}")
            continue
        seen_ordinals.add(ordinal)
        bookmark_name = f"JF_ARTICLE_{ordinal:03d}_START"
        try:
            scope, pre_title = _article_layout_scope(document, bookmark_name)
        except ValueError as exc:
            blockers.append(f"layout_adjustment_scope:{ordinal}:{exc}")
            continue

        item_report: dict[str, Any] = {
            "ordinal": ordinal,
            "bookmark": bookmark_name,
            "provenance": adjustment.get("provenance"),
        }
        if adjustment.get("trim_right_aligned_padding"):
            item_report["trimmed_right_aligned_padding_paragraphs"] = (
                _trim_right_aligned_padding(scope)
            )
        if adjustment.get("collapse_redundant_pre_title_empty_paragraphs"):
            item_report["collapsed_pre_title_empty_paragraphs"] = (
                _collapse_redundant_empty_paragraphs(document, pre_title)
            )

        style_names = [
            str(value) for value in adjustment.get("paragraph_styles_to_normal", [])
        ]
        if style_names:
            try:
                source_style_ids = {
                    _style_id_by_name(document, name) for name in style_names
                }
                normal_style_id = _style_id_by_name(document, "Normal")
            except ValueError as exc:
                blockers.append(f"layout_adjustment_style:{ordinal}:{exc}")
            else:
                item_report["paragraph_styles_normalized"] = _normalize_paragraph_styles(
                    scope,
                    source_style_ids,
                    normal_style_id,
                )
        if adjustment.get("normalize_paragraph_spacing"):
            item_report["paragraph_spacings_normalized"] = (
                _normalize_paragraph_spacing(scope)
            )
        if adjustment.get("normalize_character_spacing"):
            item_report["character_spacings_normalized"] = (
                _normalize_character_spacing(scope)
            )

        anchor = str(adjustment.get("page_break_before_anchor") or "").strip()
        if anchor:
            matches = _find_normalized_anchor(scope, anchor)
            if len(matches) != 1:
                blockers.append(
                    f"layout_page_anchor:{ordinal}:matches={len(matches)}:{anchor}"
                )
            else:
                _insert_page_break_at_anchor(*matches[0])
                item_report["page_break_inserted"] = True
                item_report["page_break_before_anchor"] = anchor
                item_report["official_physical_page"] = adjustment.get(
                    "official_physical_page"
                )
        applied.append(item_report)
    return {
        "status": "PASS" if not blockers else "BLOCKED",
        "requested": len(adjustments),
        "applied": applied,
        "blockers": blockers,
    }


def _article_layout_scope(
    document: Document,
    bookmark_name: str,
) -> tuple[list[Any], list[Any]]:
    starts = [
        item
        for item in document.element.body.xpath(".//w:bookmarkStart")
        if item.get(qn("w:name")) == bookmark_name
    ]
    if len(starts) != 1:
        raise ValueError(f"bookmark_count:{bookmark_name}:{len(starts)}")
    title = starts[0].getparent()
    if title.tag != qn("w:p") or title.getparent() is not document.element.body:
        raise ValueError(f"bookmark_not_in_body_paragraph:{bookmark_name}")
    body = document.element.body
    body_elements = list(body)
    title_index = body_elements.index(title)
    ordinal_match = re.search(r"ARTICLE_(\d+)_START$", bookmark_name)
    if ordinal_match is None:
        raise ValueError(f"invalid_article_bookmark:{bookmark_name}")
    ordinal = int(ordinal_match.group(1))

    def block_index(candidate_ordinal: int) -> int | None:
        marker_name = f"JF_LAYOUT_ARTICLE_{candidate_ordinal:03d}_START"
        markers = [
            item
            for item in body.xpath(".//w:bookmarkStart")
            if item.get(qn("w:name")) == marker_name
        ]
        if len(markers) > 1:
            raise ValueError(f"block_bookmark_count:{marker_name}:{len(markers)}")
        if not markers:
            return None
        paragraph = markers[0].getparent()
        if paragraph.tag != qn("w:p") or paragraph.getparent() is not body:
            raise ValueError(f"block_bookmark_not_in_body_paragraph:{marker_name}")
        return body_elements.index(paragraph)

    start_index = block_index(ordinal)
    if start_index is None:
        raise ValueError(f"block_bookmark_missing:{ordinal}")
    end_index = block_index(ordinal + 1)
    if end_index is None:
        end_index = next(
            (
                index
                for index in range(title_index + 1, len(body_elements))
                if body_elements[index].tag == qn("w:p")
                and body_elements[index].xpath("./w:pPr/w:sectPr")
            ),
            len(body_elements),
        )
    if not (start_index < title_index < end_index):
        raise ValueError(
            f"invalid_block_range:{bookmark_name}:{start_index}:{title_index}:{end_index}"
        )
    scope_elements = body_elements[start_index:end_index]
    scope = [
        paragraph
        for element in scope_elements
        for paragraph in element.iter(qn("w:p"))
    ]
    pre_title = body_elements[start_index + 1 : title_index]
    return scope, pre_title


def _trim_right_aligned_padding(paragraphs: list[Any]) -> int:
    count = 0
    for paragraph in paragraphs:
        alignment = paragraph.find("./" + qn("w:pPr") + "/" + qn("w:jc"))
        if alignment is None or alignment.get(qn("w:val")) not in {"right", "end"}:
            continue
        texts = list(paragraph.iter(qn("w:t")))
        full_text = "".join(item.text or "" for item in texts)
        padding = len(full_text) - len(full_text.lstrip(" \u00a0"))
        if padding < 2 or not full_text.strip():
            continue
        remaining = padding
        for text in texts:
            value = text.text or ""
            removable = min(remaining, len(value) - len(value.lstrip(" \u00a0")))
            text.text = value[removable:]
            _set_xml_space(text)
            remaining -= removable
            if remaining == 0:
                break
        count += 1
    return count


def _collapse_redundant_empty_paragraphs(
    document: Document,
    elements: list[Any],
) -> int:
    removed = 0
    previous_empty = False
    body = document.element.body
    for element in list(elements):
        if element.tag != qn("w:p"):
            previous_empty = False
            continue
        text = "".join(element.xpath(".//w:t/text()")).strip()
        protected = element.xpath(
            ".//w:tab | .//w:drawing | .//w:pict | .//w:object | .//w:instrText | "
            ".//w:fldSimple | .//w:bookmarkStart | ./w:pPr/w:sectPr"
        )
        empty = not text and not protected
        if empty and previous_empty:
            body.remove(element)
            removed += 1
            continue
        previous_empty = empty
    return removed


def _normalize_paragraph_styles(
    paragraphs: list[Any],
    source_style_ids: set[str],
    normal_style_id: str,
) -> int:
    count = 0
    for paragraph in paragraphs:
        ppr = paragraph.find(qn("w:pPr"))
        style = ppr.find(qn("w:pStyle")) if ppr is not None else None
        if style is None or style.get(qn("w:val")) not in source_style_ids:
            continue
        style.set(qn("w:val"), normal_style_id)
        count += 1
    return count


def _normalize_paragraph_spacing(paragraphs: list[Any]) -> int:
    count = 0
    for paragraph in paragraphs:
        spacing = paragraph.find("./" + qn("w:pPr") + "/" + qn("w:spacing"))
        if spacing is None:
            continue
        before = spacing.get(qn("w:before"))
        after = spacing.get(qn("w:after"))
        if before in {None, "0"} and after in {None, "0"}:
            continue
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        count += 1
    return count


def _normalize_character_spacing(paragraphs: list[Any]) -> int:
    count = 0
    for paragraph in paragraphs:
        for run in paragraph.iter(qn("w:r")):
            if not list(run.iter(qn("w:t"))):
                continue
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run.insert(0, rpr)
            spacing = rpr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                rpr.append(spacing)
            spacing.set(qn("w:val"), "0")
            count += 1
    return count


def _find_normalized_anchor(
    paragraphs: list[Any],
    anchor: str,
) -> list[tuple[Any, int]]:
    normalized_anchor = _normalized_alnum(anchor)
    matches: list[tuple[Any, int]] = []
    for paragraph in paragraphs:
        normalized, mapping = _normalized_text_mapping(paragraph)
        start = normalized.find(normalized_anchor)
        while start >= 0:
            matches.append(mapping[start])
            start = normalized.find(normalized_anchor, start + 1)
    return matches


def _normalized_text_mapping(paragraph: Any) -> tuple[str, list[tuple[Any, int]]]:
    characters: list[str] = []
    mapping: list[tuple[Any, int]] = []
    for text in paragraph.iter(qn("w:t")):
        for index, character in enumerate(text.text or ""):
            if character.isalnum():
                characters.append(character.casefold())
                mapping.append((text, index))
    return "".join(characters), mapping


def _normalized_alnum(value: str) -> str:
    return "".join(character.casefold() for character in value if character.isalnum())


def _insert_page_break_at_anchor(text: Any, offset: int) -> None:
    value = text.text or ""
    suffix = value[offset:]
    text.text = value[:offset]
    _set_xml_space(text)
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    text.addnext(page_break)
    continuation = OxmlElement("w:t")
    continuation.text = suffix
    _set_xml_space(continuation)
    page_break.addnext(continuation)


def _set_xml_space(text: Any) -> None:
    value = text.text or ""
    attribute = qn("xml:space")
    if value.startswith(" ") or value.endswith(" "):
        text.set(attribute, "preserve")
    else:
        text.attrib.pop(attribute, None)


def _style_id_by_name(document: Document, style_name: str) -> str:
    for style in document.styles:
        if style.name == style_name or style.style_id == style_name:
            return style.style_id
    raise ValueError(f"Required paragraph style was not found: {style_name}")


def _set_word_compatibility_mode(document: Document, compatibility_mode: int) -> None:
    if compatibility_mode not in range(11, 16):
        raise ValueError(f"Unsupported Word compatibility mode: {compatibility_mode}")
    compat = document.settings.element.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        document.settings.element.append(compat)
    matches = [
        item
        for item in compat.findall(qn("w:compatSetting"))
        if item.get(qn("w:name")) == "compatibilityMode"
    ]
    setting = matches[0] if matches else OxmlElement("w:compatSetting")
    if not matches:
        compat.append(setting)
    for duplicate in matches[1:]:
        compat.remove(duplicate)
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), str(compatibility_mode))


def _add_complex_field(paragraph: Any, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        paragraph.add_run()._r.append(element)
