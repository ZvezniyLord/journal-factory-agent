from __future__ import annotations

from pathlib import Path
from shutil import copyfile

from docx import Document


def build_draft(etalon: Path, output: Path, articles: list[tuple[str, str]]) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    copyfile(etalon, output)
    doc = Document(str(output))
    doc.add_page_break()
    doc.add_heading("MVP ARTICLE INVENTORY DRAFT", level=1)
    for title, text in articles:
        doc.add_page_break()
        doc.add_heading(Path(title).stem[:180], level=2)
        for paragraph in text.splitlines():
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    doc.save(str(output))
    return output

