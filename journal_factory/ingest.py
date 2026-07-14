from __future__ import annotations

from dataclasses import asdict, dataclass
from pathlib import Path
from zipfile import ZipFile
import re
import shutil
import subprocess
import tempfile

from docx import Document


ARTICLE_HINTS = (
    "стат",
    "тез",
    "тези",
    "tezy",
    "tez",
    "article",
    "thesis",
    "stattia",
    "матеріали",
    "dopovid",
)

NON_ARTICLE_HINTS = (
    "анкета",
    "anketa",
    "заяв",
    "zayav",
    "інформаційний лист",
    "informatsiinyi lyst",
    "information letter",
    "квитан",
    "receipt",
    "оплат",
    "oplata",
    "сертиф",
    "certificate",
    "шаблон",
    "template",
    ".pdf",
    ".jpg",
    ".png",
    ".jpeg",
    ".xlsx",
    ".xls",
    ".xla",
)


@dataclass
class ArchiveEntry:
    path: str
    size: int
    extension: str
    article_candidate: bool
    reason: str


def is_article_candidate(name: str) -> tuple[bool, str]:
    lower = Path(name).name.lower()
    if lower.startswith(("~$", "._")):
        return False, "office_temp_file"
    if not lower.endswith((".docx", ".doc")):
        return False, "not_word_document"
    if any(hint in lower for hint in NON_ARTICLE_HINTS):
        return False, "non_article_hint"
    if any(hint in lower for hint in ARTICLE_HINTS):
        return True, "article_hint"
    return True, "word_document_without_negative_hint"


def inventory_archive(archive: Path) -> list[ArchiveEntry]:
    entries: list[ArchiveEntry] = []
    if archive.is_dir():
        for item in archive.rglob("*"):
            if item.is_dir():
                continue
            rel = item.relative_to(archive).as_posix()
            candidate, reason = is_article_candidate(rel)
            entries.append(ArchiveEntry(rel, item.stat().st_size, item.suffix.lower(), candidate, reason))
    else:
        with ZipFile(archive) as zf:
            for item in zf.infolist():
                if item.is_dir():
                    continue
                candidate, reason = is_article_candidate(item.filename)
                entries.append(ArchiveEntry(item.filename, item.file_size, Path(item.filename).suffix.lower(), candidate, reason))
    return entries


def is_non_article_text(text: str) -> bool:
    normalized = text.lower()[:1200]
    form_markers = (
        "анкета учасника",
        "анкета-заявка",
        "заявка учасника",
        "міжнародної науково-практичної конференції",
    )
    article_markers = ("удк", "udc", "список використан", "літератур", "references")
    return any(marker in normalized for marker in form_markers) and not any(marker in normalized for marker in article_markers)


def _extract_doc_text(path: Path, limit: int) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        return ""
    try:
        result = subprocess.run([antiword, str(path)], capture_output=True, timeout=120)
    except (OSError, subprocess.SubprocessError):
        return ""
    raw = result.stdout.strip()
    text = ""
    for encoding in ("utf-8", "cp1252", "latin-1"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if not text:
        return ""
    return re.sub(r"\n{3,}", "\n\n", text)[:limit]


def extract_docx_text_from_zip(archive: Path, member: str, limit: int = 12000) -> str:
    lower = member.lower()
    if not lower.endswith((".docx", ".doc")):
        return ""
    try:
        if archive.is_dir():
            source = archive / Path(member)
            if lower.endswith(".doc"):
                return _extract_doc_text(source, limit)
            doc = Document(str(source))
        else:
            with ZipFile(archive) as zf, tempfile.TemporaryDirectory() as tmp:
                suffix = Path(member).suffix.lower()
                target = Path(tmp) / f"article{suffix}"
                target.write_bytes(zf.read(member))
                if lower.endswith(".doc"):
                    return _extract_doc_text(target, limit)
                doc = Document(str(target))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return re.sub(r"\n{3,}", "\n\n", text)[:limit]
    except Exception:
        return ""


def inventory_as_dict(entries: list[ArchiveEntry]) -> list[dict]:
    return [asdict(entry) for entry in entries]
