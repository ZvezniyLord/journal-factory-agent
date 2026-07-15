from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any
from xml.etree import ElementTree as ET

import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt
from docxcompose.composer import Composer
from PIL import Image, ImageDraw

from .config import AppConfig


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
}

STYLE_TITLE = "\u041d\u0430\u0437\u0432\u04301"
STYLE_FIGURE = "\u0420\u0418\u0421"
STYLE_FIGURE_CAPTION = "\u0420\u0438\u0441\u041f\u0456\u0434"

DOI_RE = re.compile(r"(?:https?://doi\.org/)?(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)", re.I)
ORCID_RE = re.compile(r"\b\d{4}-\d{4}-\d{4}-\d{3}[0-9X]\b")
UDC_RE = re.compile(r"\b(?:UDC|\u0423\u0414\u041a)\s*[: ]\s*([^\n\r]+)", re.I)


@dataclass(frozen=True)
class ArticleCandidate:
    article_id: str
    source_path: str
    extracted_path: Path
    source_sha256: str
    title: str
    authors: str
    has_udc: bool
    identifiers: dict[str, Any]
    table_count: int
    table_cells: list[str]
    media_hashes: list[str]
    textbox_text: list[str]
    paragraph_count: int
    visible_units: list[str]
    normalized_text: str
    score: int


def run_production_preview(config: AppConfig, limit: int = 3) -> dict[str, Any]:
    build_dir = config.build_dir
    reports_dir = build_dir / "reports"
    render_dir = build_dir / "render_v2"
    for path in (build_dir, reports_dir, render_dir):
        path.mkdir(parents=True, exist_ok=True)

    with TemporaryDirectory(prefix="journal-preview-") as tmp:
        source_root = Path(tmp) / "source"
        _materialize_source(config.archive, source_root)
        candidates = _discover_article_candidates(source_root)
        selected = _select_smoke_articles(candidates, limit)

        output_docx = build_dir / "JOURNAL_SMOKE_V2.docx"
        output_pdf = build_dir / "JOURNAL_SMOKE_V2.pdf"
        if not selected:
            audit = {
                "status": "FAIL",
                "blockers": ["NO_ARTICLES_SELECTED"],
                "articles": [],
            }
            _write_json(reports_dir / "SMOKE_AUDIT_V2.json", audit)
            return {"status": "REVIEW", "articles": 0, "audit": str(reports_dir / "SMOKE_AUDIT_V2.json")}

        gemma_decisions = _collect_udc_decisions(selected)
        frontmatter_audit = _assemble_smoke_docx(config.etalon, output_docx, selected)
        _render_pdf_and_pages(output_docx, output_pdf, render_dir)
        page_numbers = _detect_article_pages(output_pdf, selected)
        _update_toc_page_numbers(output_docx, page_numbers)
        _render_pdf_and_pages(output_docx, output_pdf, render_dir)
        contact_sheet = _make_contact_sheet(render_dir)

        final_info = _inspect_docx(output_docx)
        identifier_audit = [_identifier_audit(item, final_info) for item in selected]
        article_audits = [_audit_article(item, final_info, page_numbers.get(item.article_id)) for item in selected]
        toc_audit = _export_toc(output_docx, reports_dir / "toc_export.tsv")
        visible_defects = _collect_defects(frontmatter_audit, identifier_audit, article_audits, toc_audit, output_pdf, contact_sheet)

        smoke_audit = {
            "status": "FAIL" if _has_fail_defects(visible_defects) else "REVIEW",
            "pass_forbidden": True,
            "mode": "production-preview",
            "docx": str(output_docx),
            "pdf": str(output_pdf),
            "render_dir": str(render_dir),
            "contact_sheet": str(contact_sheet),
            "article_count": len(selected),
            "frontmatter_metadata_audit": str(reports_dir / "frontmatter_metadata_audit.json"),
            "identifier_audit": str(reports_dir / "identifier_audit.json"),
            "toc_export": str(reports_dir / "toc_export.tsv"),
            "toc_audit": toc_audit,
            "articles": article_audits,
            "identifiers": identifier_audit,
            "gemma": gemma_decisions,
            "gemma_note": "UDC proposal requested only when a selected article lacks literal UDC."
            if gemma_decisions
            else "No selected real article lacked literal UDC; no UDC proposal was requested.",
            "visible_defects": visible_defects,
            "unknown_text_deletion_count": sum(len(item["deleted_text"]) for item in article_audits),
            "unknown_object_deletion_count": sum(1 for item in article_audits if item["lost_tables"] or item["lost_media"] or item["lost_textbox_text"]),
            "final_paragraph_count": final_info["paragraph_count"],
            "final_table_count": final_info["table_count"],
        }
        _write_json(reports_dir / "frontmatter_metadata_audit.json", frontmatter_audit)
        _write_json(reports_dir / "identifier_audit.json", {"status": _status(identifier_audit), "articles": identifier_audit})
        _write_json(reports_dir / "SMOKE_AUDIT_V2.json", smoke_audit)
        return {
            "status": "REVIEW",
            "articles": len(selected),
            "docx": str(output_docx),
            "pdf": str(output_pdf),
            "contact_sheet": str(contact_sheet),
            "audit": str(reports_dir / "SMOKE_AUDIT_V2.json"),
            "visible_defects": visible_defects,
        }


def _materialize_source(source: Path, destination: Path) -> None:
    if destination.exists():
        shutil.rmtree(destination)
    destination.mkdir(parents=True)
    if source.is_dir():
        shutil.copytree(source, destination, dirs_exist_ok=True)
        return
    if not zipfile.is_zipfile(source):
        raise ValueError(f"Source must be a ZIP or directory: {source}")
    with zipfile.ZipFile(source) as archive:
        root = destination.resolve()
        for member in archive.infolist():
            target = (destination / member.filename).resolve()
            if root not in target.parents and target != root:
                raise ValueError(f"Blocked unsafe ZIP member: {member.filename}")
            archive.extract(member, destination)


def _discover_article_candidates(source_root: Path) -> list[ArticleCandidate]:
    candidates: list[ArticleCandidate] = []
    for path in sorted(source_root.rglob("*.docx")):
        if _is_service_file(path):
            continue
        lower = path.name.lower()
        negative_tokens = (
            "\u0430\u043d\u043a\u0435\u0442\u0430",
            "anketa",
            "\u0437\u0430\u044f\u0432",
            "\u043e\u043f\u043b\u0430\u0442",
            "oplata",
            "\u043a\u0432\u0438\u0442\u0430\u043d",
            "\u0441\u0435\u0440\u0442\u0438\u0444",
            "\u0456\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0456\u0439\u043d",
        )
        if any(token in lower for token in negative_tokens):
            continue
        try:
            info = _inspect_docx(path)
        except Exception:
            continue
        if not info["normalized_text"]:
            continue
        score = (
            (20 if info["identifiers"]["doi"] else 0)
            + min(len(info["identifiers"]["orcids"]), 3) * 10
            + info["table_count"] * 5
            + len(info["media_hashes"]) * 5
            + (5 if info["long_frontmatter"] else 0)
            + (3 if info["reference_count"] else 0)
            + (4 if not info["has_udc"] else 0)
        )
        candidates.append(
            ArticleCandidate(
                article_id=f"article-{len(candidates) + 1:03d}",
                source_path=_repair_mojibake(path.relative_to(source_root).as_posix()),
                extracted_path=path,
                source_sha256=_sha256_file(path),
                title=info["title"],
                authors=info["authors"],
                has_udc=info["has_udc"],
                identifiers=info["identifiers"],
                table_count=info["table_count"],
                table_cells=info["table_cells"],
                media_hashes=info["media_hashes"],
                textbox_text=info["textbox_text"],
                paragraph_count=info["paragraph_count"],
                visible_units=info["visible_units"],
                normalized_text=info["normalized_text"],
                score=score,
            )
        )
    return candidates


def _select_smoke_articles(candidates: list[ArticleCandidate], limit: int) -> list[ArticleCandidate]:
    selected: list[ArticleCandidate] = []

    def add(predicate) -> None:
        if len(selected) >= limit:
            return
        for candidate in sorted(candidates, key=lambda item: item.score, reverse=True):
            if candidate not in selected and predicate(candidate):
                selected.append(candidate)
                return

    add(lambda item: bool(item.identifiers["doi"]) and len(item.identifiers["orcids"]) >= 2 and (item.table_count or item.media_hashes))
    add(lambda item: bool(item.identifiers["doi"]) and len(item.identifiers["orcids"]) >= 1)
    add(lambda item: bool(item.identifiers["doi"]) and (item.table_count or item.media_hashes))
    add(lambda item: item.identifiers["doi"] and item.identifiers["udc"])
    for candidate in sorted(candidates, key=lambda item: item.score, reverse=True):
        if len(selected) >= limit:
            break
        if candidate not in selected:
            selected.append(candidate)
    return selected[:limit]


def _assemble_smoke_docx(etalon: Path, output_docx: Path, selected: list[ArticleCandidate]) -> dict[str, Any]:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(etalon, output_docx)
    doc = Document(str(output_docx))
    _ensure_preview_styles(doc)
    replacements = _frontmatter_replacements()
    replacement_hits = _replace_frontmatter_text(doc, replacements)
    _insert_canonical_toc(doc, selected)
    doc.save(str(output_docx))

    master = Document(str(output_docx))
    composer = Composer(master)
    for candidate in selected:
        page_break_doc = Document()
        page_break_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        composer.append(page_break_doc)
        composer.append(Document(str(candidate.extracted_path)))
    composer.save(str(output_docx))
    _apply_smoke_styles(output_docx)
    return _audit_frontmatter_metadata(output_docx, replacement_hits)


def _frontmatter_replacements() -> dict[str, str]:
    return {
        "JUNE 12-14, 2026": "JULY 6-8, 2026",
        "NEW YORK, USA": "KHARKIV, UKRAINE",
        "January 19-21, 2026": "July 6-8, 2026",
        "Cambridge, United Kingdom": "Kharkiv, Ukraine",
        "February 8, 2026": "July 15, 2026",
        "https://naukainfo.com/conference?id=91": "https://naukainfo.com/conference?id=138",
        "Proceedings of the International scientific and practical conference \u201cScience in the Modern World\u201d (January 19-21, 2026)": "Proceedings of the International scientific and practical conference \u201cScience in the Modern World\u201d (July 6-8, 2026)",
        "ISBN 978-617-8680-36-7": "ISBN pending operator review",
        "https://doi.org/10.64828/conf-91-2026": "https://doi.org/10.64828/conf-138-2026",
    }


def _replace_frontmatter_text(doc: Document, replacements: dict[str, str]) -> dict[str, int]:
    hits = {key: 0 for key in replacements}
    for paragraph in doc.paragraphs[:90]:
        original = paragraph.text
        updated = original
        for old, new in replacements.items():
            if old in updated:
                hits[old] += updated.count(old)
                updated = updated.replace(old, new)
        if updated != original:
            _set_paragraph_text_preserve_element(paragraph, updated)
    return hits


def _set_paragraph_text_preserve_element(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _insert_canonical_toc(doc: Document, selected: list[ArticleCandidate]) -> None:
    target = _find_toc_anchor(doc)
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    rows: list[list[str]] = [["", "SCIENCE IN THE MODERN WORLD", ""]]
    for index, candidate in enumerate(selected, start=1):
        rows.append([str(index), candidate.authors, ""])
        rows.append(["", candidate.title, ""])
    for values in rows:
        cells = table.add_row().cells
        style = "TabSEC" if values[0] == "" and values[2] == "" and values[1] == "SCIENCE IN THE MODERN WORLD" else "TabTaitl" if values[0] == "" else "TabPIP"
        for i, value in enumerate(values):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                _clear_paragraph_layout(paragraph)
                _safe_style(paragraph, style)
    if target is not None:
        target._p.addnext(table._tbl)


def _find_toc_anchor(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().upper() == "TABLE OF CONTENTS":
            return paragraph
    return doc.paragraphs[-1] if doc.paragraphs else None


def _ensure_preview_styles(doc: Document) -> None:
    for style_name in ("TabSEC", "TabPIP", "TabTaitl"):
        if style_name not in [style.name for style in doc.styles]:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Times New Roman"
            style.font.size = Pt(10)


def _apply_smoke_styles(output_docx: Path) -> None:
    doc = Document(str(output_docx))
    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        _clear_paragraph_layout(paragraph)
        if _is_doi(text):
            _safe_style(paragraph, "pip")
            in_refs = False
        elif _is_udc(text):
            _safe_style(paragraph, "UDC")
            in_refs = False
        elif _looks_like_references_title(text):
            _safe_style(paragraph, "REF-TITLE")
            in_refs = True
        elif in_refs:
            _safe_style(paragraph, "REFER")
        elif _looks_like_figure_caption(text):
            _safe_style(paragraph, STYLE_FIGURE_CAPTION)
        elif _looks_like_table_caption(text):
            _safe_style(paragraph, STYLE_FIGURE)
        elif _looks_like_title(text):
            _safe_style(paragraph, STYLE_TITLE)
            in_refs = False
        elif _looks_like_author_line(text):
            _safe_style(paragraph, "AUTOR")
            in_refs = False
        else:
            _safe_style(paragraph, "pip")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _clear_paragraph_layout(paragraph)
                    if paragraph.style.name not in {"TabSEC", "TabPIP", "TabTaitl"}:
                        _safe_style(paragraph, "TABLETEXT")
    doc.save(str(output_docx))
    _remove_unnecessary_breaks_and_tabs(output_docx)


def _clear_paragraph_layout(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = None
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.tab_stops.clear_all()


def _remove_unnecessary_breaks_and_tabs(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    root = ET.fromstring(files["word/document.xml"])
    for para in root.findall(".//w:p", NS):
        texts = [node.text or "" for node in para.findall(".//w:t", NS)]
        joined = "".join(texts)
        if _is_frontmatter_like(joined):
            for br in list(para.findall(".//w:br", NS)):
                parent = _find_parent(root, br)
                if parent is not None:
                    parent.remove(br)
            for tab in list(para.findall(".//w:tab", NS)):
                parent = _find_parent(root, tab)
                if parent is not None:
                    parent.remove(tab)
    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, payload in files.items():
            zout.writestr(name, payload)
    tmp.replace(path)


def _find_parent(root: ET.Element, child: ET.Element) -> ET.Element | None:
    for parent in root.iter():
        if child in list(parent):
            return parent
    return None


def _is_frontmatter_like(text: str) -> bool:
    return any(token in text for token in ("ORCID", "orcid.org", "@", "+", "\u043a\u0430\u0444\u0435\u0434\u0440", "\u0443\u043d\u0456\u0432\u0435\u0440\u0441\u0438\u0442"))


def _inspect_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [_repair_mojibake(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    table_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = _normalize_text(_repair_mojibake(cell.text))
                if text:
                    table_cells.append(text)
    media_hashes = _media_hashes(path)
    textbox_text = _textbox_text(path)
    visible_units = [_normalize_text(item) for item in paragraphs + table_cells + textbox_text if _normalize_text(item)]
    full_text = "\n".join(visible_units)
    identifiers = _extract_identifiers(full_text)
    return {
        "title": _detect_title(paragraphs, path),
        "authors": _detect_authors(paragraphs),
        "has_udc": bool(identifiers["udc"]),
        "identifiers": identifiers,
        "table_count": len(doc.tables),
        "table_cells": table_cells,
        "media_hashes": media_hashes,
        "textbox_text": textbox_text,
        "reference_count": _count_references(paragraphs),
        "paragraph_count": len(paragraphs),
        "style_ids": _collect_style_ids(path),
        "visible_units": visible_units,
        "normalized_text": _normalize_text(full_text),
        "long_frontmatter": any(len(item) > 120 for item in paragraphs[:14]),
    }


def _media_hashes(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as package:
        return sorted(
            hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/")
        )


def _textbox_text(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as package:
        if "word/document.xml" not in package.namelist():
            return []
        root = ET.fromstring(package.read("word/document.xml"))
    values = []
    for txbx in root.findall(".//w:txbxContent", NS):
        text = " ".join(node.text or "" for node in txbx.findall(".//w:t", NS))
        text = _normalize_text(_repair_mojibake(text))
        if text:
            values.append(text)
    return values


def _extract_identifiers(text: str) -> dict[str, Any]:
    dois = []
    for value in DOI_RE.findall(text):
        clean = value.rstrip(".,;)")
        if clean not in dois:
            dois.append(clean)
    orcids = []
    for value in ORCID_RE.findall(text):
        if value not in orcids:
            orcids.append(value)
    udcs = []
    for value in UDC_RE.findall(text):
        clean = _normalize_text(value).strip(" .;")
        if clean and clean not in udcs:
            udcs.append(clean)
    return {"doi": dois, "udc": udcs, "orcids": orcids}


def _identifier_audit(candidate: ArticleCandidate, final_info: dict[str, Any]) -> dict[str, Any]:
    final_text = final_info["normalized_text"]
    source = candidate.identifiers
    final_doi = [doi for doi in source["doi"] if doi in final_text]
    final_orcids = [orcid for orcid in source["orcids"] if orcid in final_text]
    final_udc = [udc for udc in source["udc"] if udc in final_text]
    missing_doi = [doi for doi in source["doi"] if doi not in final_doi]
    missing_orcids = [orcid for orcid in source["orcids"] if orcid not in final_orcids]
    missing_udc = [udc for udc in source["udc"] if udc not in final_udc]
    source_requires_doi_before_udc = _doi_before_udc(candidate.normalized_text, source["doi"], source["udc"])
    doi_before_udc = True if not source_requires_doi_before_udc else _doi_before_udc(final_text, source["doi"], source["udc"])
    blockers = []
    if missing_doi or missing_orcids:
        blockers.append("REQUIRED_IDENTIFIER_LOST")
    if not doi_before_udc and source["doi"] and source["udc"]:
        blockers.append("IDENTIFIER_ORDER_INVALID")
    return {
        "article_id": candidate.article_id,
        "source_path": candidate.source_path,
        "source_doi": source["doi"],
        "final_doi": final_doi,
        "source_udc": source["udc"],
        "final_udc": final_udc,
        "source_orcids": source["orcids"],
        "final_orcids": final_orcids,
        "missing_doi": missing_doi,
        "missing_orcids": missing_orcids,
        "missing_udc": missing_udc,
        "doi_before_udc": doi_before_udc,
        "source_requires_doi_before_udc": source_requires_doi_before_udc,
        "blockers": blockers,
        "status": "FAIL" if blockers else "REVIEW",
    }


def _audit_article(candidate: ArticleCandidate, final_info: dict[str, Any], start_page: int | None) -> dict[str, Any]:
    final_text = final_info["normalized_text"]
    deleted_text = [unit for unit in candidate.visible_units if unit and unit not in final_text]
    source_media = set(candidate.media_hashes)
    final_media = set(final_info["media_hashes"])
    lost_media = sorted(source_media - final_media)
    lost_textbox = [text for text in candidate.textbox_text if text not in final_text]
    lost_tables = candidate.table_count > final_info["table_count"]
    blockers = []
    if deleted_text:
        blockers.append("SOURCE_TEXT_LOSS")
    if lost_tables or lost_media or lost_textbox:
        blockers.append("SOURCE_OBJECT_LOSS")
    return {
        "article_id": candidate.article_id,
        "journal_order": start_page,
        "source_path": candidate.source_path,
        "source_sha256": candidate.source_sha256,
        "title": candidate.title,
        "authors": candidate.authors,
        "source_meaningful_paragraphs": candidate.paragraph_count,
        "source_visible_unit_count": len(candidate.visible_units),
        "source_tables": candidate.table_count,
        "source_table_cells": candidate.table_cells,
        "source_media_count": len(candidate.media_hashes),
        "source_media_sha256": candidate.media_hashes,
        "source_textbox_text": candidate.textbox_text,
        "deleted_text": deleted_text,
        "inserted_text": [],
        "replaced_text": [],
        "missing_doi": [],
        "missing_orcids": [],
        "missing_udc": [],
        "lost_tables": lost_tables,
        "lost_media": lost_media,
        "lost_textbox_text": lost_textbox,
        "actual_style_ids": final_info["style_ids"],
        "starts_on_new_page": start_page is not None,
        "warnings": ["business_normalization_incomplete_review"],
        "blockers": blockers,
        "status": "FAIL" if blockers else "REVIEW",
    }


def _doi_before_udc(final_text: str, dois: list[str], udcs: list[str]) -> bool:
    if not dois or not udcs:
        return True
    doi_positions = [final_text.find(doi) for doi in dois if final_text.find(doi) >= 0]
    udc_positions = [final_text.find(udc) for udc in udcs if final_text.find(udc) >= 0]
    return bool(doi_positions and udc_positions and min(doi_positions) < min(udc_positions))


def _audit_frontmatter_metadata(output_docx: Path, replacement_hits: dict[str, int]) -> dict[str, Any]:
    doc = Document(str(output_docx))
    front_text = "\n".join(p.text for p in doc.paragraphs[:90])
    forbidden = [
        "Cambridge, United Kingdom",
        "NEW YORK, USA",
        "January 19-21, 2026",
        "JUNE 12-14, 2026",
        "conference?id=91",
        "conf-91-2026",
        "978-617-8680-36-7",
    ]
    remaining = [item for item in forbidden if item in front_text]
    required = ["KHARKIV, UKRAINE", "July 6-8, 2026", "conference?id=138", "conf-138-2026"]
    missing = [item for item in required if item not in front_text]
    blockers = []
    if remaining or missing:
        blockers.append("FRONTMATTER_METADATA_MISMATCH")
    return {
        "status": "FAIL" if blockers else "REVIEW",
        "replacement_hits": replacement_hits,
        "forbidden_remaining": remaining,
        "required_missing": missing,
        "institutional_header_present": "INFORMATION PLATFORM" in front_text and "UKRAINIAN INSTITUTE" in front_text,
        "blockers": blockers,
    }


def _export_toc(output_docx: Path, tsv_path: Path) -> dict[str, Any]:
    doc = Document(str(output_docx))
    rows = []
    toc_table = None
    for table in doc.tables:
        if table.rows and len(table.rows[0].cells) == 3 and "SCIENCE IN THE MODERN WORLD" in table.rows[0].cells[1].text:
            toc_table = table
            break
    if toc_table is not None:
        for row in toc_table.rows:
            rows.append([_normalize_text(cell.text) for cell in row.cells])
    with tsv_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter="\t")
        writer.writerows(rows)
    blockers = []
    if toc_table is None or any(len(row) != 3 for row in rows):
        blockers.append("TOC_CONTRACT_INVALID")
    if any(row[0] and row[1].startswith(row[0] + ".") for row in rows):
        blockers.append("TOC_CONTRACT_INVALID")
    return {
        "status": "FAIL" if blockers else "REVIEW",
        "row_count": len(rows),
        "column_count": 3 if rows else 0,
        "rows": rows,
        "decorative_rows": [],
        "blockers": sorted(set(blockers)),
    }


def _update_toc_page_numbers(output_docx: Path, page_numbers: dict[str, int]) -> None:
    doc = Document(str(output_docx))
    article_ids = list(page_numbers)
    idx = 0
    for table in doc.tables:
        if table.rows and len(table.rows[0].cells) == 3 and "SCIENCE IN THE MODERN WORLD" in table.rows[0].cells[1].text:
            for row in table.rows[1:]:
                if row.cells[0].text.strip() and idx < len(article_ids):
                    row.cells[2].text = str(page_numbers[article_ids[idx]])
                    idx += 1
            break
    doc.save(str(output_docx))


def _detect_article_pages(pdf: Path, selected: list[ArticleCandidate]) -> dict[str, int]:
    if not pdf.exists() or not shutil.which("pdftotext"):
        return {}
    pages: dict[int, str] = {}
    total = _pdf_page_count(pdf)
    for page in range(1, total + 1):
        result = subprocess.run(
            ["pdftotext", "-f", str(page), "-l", str(page), str(pdf), "-"],
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            text=True,
            timeout=30,
            check=False,
        )
        pages[page] = _normalize_text(result.stdout)
    detected = {}
    for candidate in selected:
        service_dois = [doi for doi in candidate.identifiers["doi"] if "conf-138" in doi]
        needles = service_dois[:1] or candidate.identifiers["doi"][:1] or candidate.identifiers["orcids"][:1]
        if not needles:
            needles = [candidate.title[:40]]
        for page, text in pages.items():
            if page < 6:
                continue
            if any(needle and needle in text for needle in needles):
                detected[candidate.article_id] = page
                break
    return detected


def _pdf_page_count(pdf: Path) -> int:
    if not shutil.which("pdfinfo"):
        return 0
    result = subprocess.run(["pdfinfo", str(pdf)], stdout=subprocess.PIPE, stderr=subprocess.DEVNULL, text=True, check=False)
    match = re.search(r"Pages:\s+(\d+)", result.stdout)
    return int(match.group(1)) if match else 0


def _render_pdf_and_pages(output_docx: Path, output_pdf: Path, render_dir: Path) -> None:
    render_dir.mkdir(parents=True, exist_ok=True)
    for old in list(render_dir.glob("page-*.png")) + list(render_dir.glob("contact_sheet.png")):
        old.unlink()
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_pdf.parent), str(output_docx)],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        timeout=180,
        check=False,
    )
    generated_pdf = output_docx.with_suffix(".pdf")
    if generated_pdf.exists() and generated_pdf != output_pdf:
        shutil.move(str(generated_pdf), str(output_pdf))
    pdftoppm = shutil.which("pdftoppm")
    if output_pdf.exists() and pdftoppm:
        subprocess.run(
            [pdftoppm, "-png", "-r", "120", str(output_pdf), str(render_dir / "page")],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            timeout=180,
            check=False,
        )


def _make_contact_sheet(render_dir: Path) -> Path:
    pages = sorted(render_dir.glob("page-*.png"))
    output = render_dir / "contact_sheet.png"
    if not pages:
        return output
    thumbs = []
    for page in pages:
        image = Image.open(page).convert("RGB")
        image.thumbnail((360, 510))
        thumbs.append((page.name, image.copy()))
    width = 760
    cell_w = 380
    cell_h = 560
    rows = (len(thumbs) + 1) // 2
    sheet = Image.new("RGB", (width, max(1, rows) * cell_h), "white")
    draw = ImageDraw.Draw(sheet)
    for i, (name, image) in enumerate(thumbs):
        x = (i % 2) * cell_w + 10
        y = (i // 2) * cell_h + 30
        draw.text((x, y - 24), name, fill="black")
        sheet.paste(image, (x, y))
    sheet.save(output)
    return output


def _collect_udc_decisions(selected: list[ArticleCandidate]) -> list[dict[str, Any]]:
    return [_ask_gemma_udc(candidate) for candidate in selected if not candidate.has_udc]


def _ask_gemma_udc(candidate: ArticleCandidate) -> dict[str, Any]:
    base_url = os.environ.get("OLLAMA_BASE_URL", "http://host.docker.internal:11434").rstrip("/")
    model = "gemma2:2b"
    expected = {
        "decision_type": "udc_proposal",
        "article_id": candidate.article_id,
        "status": "needs_operator_review",
        "value": "",
        "confidence": 0.0,
        "evidence": [],
        "model": model,
        "prompt_version": "udc-smoke-v1",
        "source_hash": candidate.source_sha256,
    }
    try:
        response = requests.post(
            f"{base_url}/api/generate",
            json={"model": model, "prompt": json.dumps(expected, ensure_ascii=False), "stream": False, "format": "json"},
            timeout=120,
        )
        response.raise_for_status()
        raw = response.json().get("response", "")
        parsed = json.loads(raw)
    except Exception as exc:
        return {**expected, "status": "model_unavailable", "raw_response": "", "valid_json": False, "error": str(exc)}
    valid = isinstance(parsed, dict) and all(parsed.get(key) == value for key, value in expected.items() if key != "value")
    parsed["raw_response"] = raw
    parsed["valid_json"] = bool(valid)
    return parsed


def _collect_defects(
    frontmatter_audit: dict[str, Any],
    identifier_audit: list[dict[str, Any]],
    article_audits: list[dict[str, Any]],
    toc_audit: dict[str, Any],
    output_pdf: Path,
    contact_sheet: Path,
) -> list[str]:
    defects = []
    defects.extend(frontmatter_audit["blockers"])
    defects.extend(toc_audit["blockers"])
    for item in identifier_audit:
        defects.extend(item["blockers"])
    for item in article_audits:
        defects.extend(item["blockers"])
        defects.extend(item["warnings"])
    if not output_pdf.exists():
        defects.append("PDF_RENDER_MISSING")
    if not contact_sheet.exists():
        defects.append("CONTACT_SHEET_MISSING")
    return sorted(set(defects))


def _has_fail_defects(defects: list[str]) -> bool:
    return any(
        item
        in {
            "FRONTMATTER_METADATA_MISMATCH",
            "TOC_CONTRACT_INVALID",
            "REQUIRED_IDENTIFIER_LOST",
            "SOURCE_TEXT_LOSS",
            "SOURCE_OBJECT_LOSS",
            "PDF_RENDER_MISSING",
            "CONTACT_SHEET_MISSING",
        }
        for item in defects
    )


def _status(items: list[dict[str, Any]]) -> str:
    return "FAIL" if any(item["status"] == "FAIL" for item in items) else "REVIEW"


def _detect_title(paragraphs: list[str], path: Path) -> str:
    for text in paragraphs[:18]:
        if _is_doi(text) or _is_udc(text) or _looks_like_author_line(text) or "ORCID" in text:
            continue
        if len(text) > 24:
            return text[:180]
    return _repair_mojibake(path.stem[:180])


def _detect_authors(paragraphs: list[str]) -> str:
    values = []
    for text in paragraphs[:12]:
        if _is_doi(text) or _is_udc(text) or "ORCID" in text or len(text) > 140:
            continue
        if _looks_like_author_line(text) or len(values) < 2:
            values.append(text)
    return "; ".join(values[:3]) or "AUTHOR REVIEW"


def _count_references(paragraphs: list[str]) -> int:
    in_refs = False
    count = 0
    for paragraph in paragraphs:
        if _looks_like_references_title(paragraph):
            in_refs = True
            continue
        if in_refs and (re.match(r"^\s*\d+[\.\)]\s+", paragraph) or len(paragraph) > 30):
            count += 1
    return count


def _collect_style_ids(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as package:
        root = ET.fromstring(package.read("word/document.xml"))
    counts: dict[str, int] = {}
    for p_style in root.findall(".//w:pStyle", NS):
        style_id = p_style.attrib.get(f"{{{NS['w']}}}val")
        if style_id:
            counts[style_id] = counts.get(style_id, 0) + 1
    return counts


def _safe_style(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except Exception:
        return


def _is_doi(text: str) -> bool:
    return bool(DOI_RE.search(text))


def _is_udc(text: str) -> bool:
    return bool(UDC_RE.search(text))


def _looks_like_author_line(text: str) -> bool:
    return len(text) < 180 and bool(re.search(r"\b[A-ZА-ЯІЇЄҐ][a-zа-яіїєґ'’.-]+\s+[A-ZА-ЯІЇЄҐ]\.", text))


def _looks_like_title(text: str) -> bool:
    if len(text) < 35 or _is_doi(text) or _is_udc(text):
        return False
    upper_chars = [ch for ch in text if ch.isalpha()]
    return bool(upper_chars) and sum(ch.isupper() for ch in upper_chars) / len(upper_chars) > 0.65


def _looks_like_references_title(text: str) -> bool:
    return text.strip().lower() in {
        "\u0441\u043f\u0438\u0441\u043e\u043a \u0432\u0438\u043a\u043e\u0440\u0438\u0441\u0442\u0430\u043d\u0438\u0445 \u0434\u0436\u0435\u0440\u0435\u043b",
        "\u043b\u0456\u0442\u0435\u0440\u0430\u0442\u0443\u0440\u0430",
        "references",
        "reference",
    }


def _looks_like_figure_caption(text: str) -> bool:
    return bool(re.match(r"^(рис\.?|figure)\s*\d+", text.strip(), re.IGNORECASE))


def _looks_like_table_caption(text: str) -> bool:
    return bool(re.match(r"^(табл\.?|таблиця|table)\s*\d+", text.strip(), re.IGNORECASE))


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _repair_mojibake(text: str) -> str:
    if not any(marker in text for marker in ("Ð", "Ã", "Рќ", "Р°", "СЏ", "С–", "С€")):
        return text
    for encoding in ("cp1251", "latin1"):
        try:
            repaired = text.encode(encoding).decode("utf-8")
        except UnicodeError:
            continue
        if repaired.count("�") <= text.count("�") and ("Р" not in repaired or repaired.count("Р") < text.count("Р")):
            return repaired
    return text


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _is_service_file(path: Path) -> bool:
    return path.name.startswith("~$") or path.name.startswith("._")


def _write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
