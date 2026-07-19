from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from journal_factory.style_bridge import merge_template_styles


def test_missing_required_styles_are_imported_by_name(tmp_path: Path) -> None:
    etalon_path = tmp_path / "etalon.docx"
    template_path = tmp_path / "Jurnal.dotx"
    output_path = tmp_path / "styled-master.docx"

    etalon = Document()
    etalon.add_paragraph("TABLE OF CONTENTS")
    etalon.save(etalon_path)

    template = Document()
    for name in ("AUTOR", "SECTION", "Назва1"):
        template.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    template.save(template_path)

    report = merge_template_styles(etalon_path, template_path, output_path)

    assert report["status"] == "PASS"
    assert report["missing_required_style_names"] == []
    style_names = {style.name for style in Document(output_path).styles}
    assert {"AUTOR", "SECTION", "Назва1"}.issubset(style_names)


def test_missing_style_blocks_build(tmp_path: Path) -> None:
    etalon_path = tmp_path / "etalon.docx"
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "styled-master.docx"

    Document().save(etalon_path)
    Document().save(template_path)

    report = merge_template_styles(etalon_path, template_path, output_path)

    assert report["status"] == "BLOCKED"
    assert set(report["missing_required_style_names"]) == {"AUTOR", "SECTION", "Назва1"}
