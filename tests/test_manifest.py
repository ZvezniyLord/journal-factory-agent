from pathlib import Path
import json
import shutil

from docx import Document

from journal_factory.archive_workspace import prepare_archive_workspace
from journal_factory.manifest import build_article_manifest


FIXTURE = Path("fixtures/manifest")


def _workspace_from_fixture(tmp_path: Path) -> Path:
    source = tmp_path / "fixture"
    shutil.copytree(FIXTURE, source)
    report = prepare_archive_workspace(source, tmp_path / "build" / "workspace", tmp_path / "build" / "reports")
    return Path(report["workspace_source"])


def _load_manifest(path: Path) -> dict:
    manifest = path / "applications" / "manifest.json"
    return json.loads(manifest.read_text(encoding="utf-8"))


def _write_manifest(path: Path, data: dict) -> None:
    (path / "applications" / "manifest.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def test_article_matched_by_author_and_title(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    assert manifest["manifest_status"] == "PASS"
    assert [item["match_status"] for item in manifest["articles"]] == ["MATCHED", "MATCHED"]


def test_filename_only_match_rejected(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)
    data = _load_manifest(workspace)
    data["articles"][0]["authors_manifest"] = ["Nobody"]
    data["articles"][0]["title_manifest"] = "article_alpha"
    _write_manifest(workspace, data)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    assert manifest["articles"][0]["match_status"] == "BLOCKED"


def test_author_only_match_rejected(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)
    data = _load_manifest(workspace)
    data["articles"][0]["title_manifest"] = "Missing title"
    _write_manifest(workspace, data)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    assert "title_evidence_missing" in manifest["articles"][0]["blockers"]


def test_title_only_match_rejected(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)
    data = _load_manifest(workspace)
    data["articles"][0]["authors_manifest"] = ["Missing Author"]
    _write_manifest(workspace, data)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    assert "author_evidence_missing" in manifest["articles"][0]["blockers"]


def test_duplicate_article_match_blocked(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)
    data = _load_manifest(workspace)
    duplicate = dict(data["articles"][0])
    duplicate["article_id"] = "synthetic-duplicate"
    duplicate["journal_order"] = 3
    data["articles"].append(duplicate)
    _write_manifest(workspace, data)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    duplicate_entries = [item for item in manifest["articles"] if item["source_path"] == "articles/article_alpha.docx"]
    assert all("source_article_matches_multiple_manifest_records" in item["blockers"] for item in duplicate_entries)


def test_duplicate_journal_order_blocked(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)
    data = _load_manifest(workspace)
    data["articles"][1]["journal_order"] = data["articles"][0]["journal_order"]
    _write_manifest(workspace, data)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    assert "duplicate_journal_order:1" in manifest["blockers"]


def test_free_listener_excluded_from_articles(tmp_path: Path) -> None:
    workspace = _workspace_from_fixture(tmp_path)

    manifest = build_article_manifest(workspace, tmp_path / "build" / "reports")

    assert manifest["article_count"] == 2
    assert manifest["free_listener_count"] == 1
    assert manifest["free_listeners"][0]["full_name"] == "Free Listener"


def test_auto_manifest_classifies_services_duplicates_and_folder_order(tmp_path: Path) -> None:
    workspace = tmp_path / "raw"
    first = workspace / "Заявки" / "1 Alpha"
    third = workspace / "Заявки" / "3 Gamma"
    first.mkdir(parents=True)
    third.mkdir(parents=True)

    _write_article(first / "article.docx", "Alpha Author", "ALPHA ARTICLE TITLE")
    _write_application(first / "application.docx", "Alpha Author", "ALPHA ARTICLE TITLE", "3. Test Section")
    (first / "receipt.pdf").write_bytes(b"%PDF-1.4 synthetic receipt")
    _write_article(third / "article.docx", "Gamma Author", "GAMMA ARTICLE TITLE")
    _write_article(
        third / "New_article.docx",
        "Gamma Author",
        "GAMMA ARTICLE TITLE",
        udc="УДК [ВКАЖІТЬ_КОД]",
    )
    _write_article(workspace / "Збірник" / "compiled.docx", "Compiled Author", "COMPILED JOURNAL TITLE")

    reports = tmp_path / "reports"
    manifest = build_article_manifest(workspace, reports)
    evidence = json.loads((reports / "manifest_evidence.json").read_text(encoding="utf-8"))

    assert manifest["manifest_status"] == "PASS"
    assert manifest["generation_mode"] == "AUTO_INVENTORY_V1"
    assert [entry["provenance"]["participant_order"] for entry in manifest["articles"]] == [1, 3]
    assert all(entry["confidence"] > 0 for entry in manifest["articles"])
    assert all(entry["provenance"]["generation"] == "AUTO_INVENTORY_V1" for entry in manifest["articles"])
    assert evidence["classification_counts"]["ARTICLE"] == 2
    assert evidence["classification_counts"]["DUPLICATE"] == 1
    assert evidence["classification_counts"]["NON_ARTICLE"] == 3
    selected_third = manifest["articles"][1]["source_path"]
    assert selected_third.endswith("/article.docx")


def test_auto_manifest_blocks_non_equivalent_articles_in_one_participant_folder(tmp_path: Path) -> None:
    folder = tmp_path / "raw" / "applications" / "1 Alpha"
    folder.mkdir(parents=True)
    _write_article(folder / "first.docx", "Alpha Author", "FIRST INDEPENDENT ARTICLE")
    _write_article(folder / "second.docx", "Alpha Author", "SECOND UNRELATED ARTICLE", body_token="different")

    manifest = build_article_manifest(tmp_path / "raw", tmp_path / "reports")

    assert manifest["manifest_status"] == "BLOCKED"
    assert any("multiple_non_equivalent_article_candidates" in item for item in manifest["blockers"])


def test_auto_manifest_marks_embedded_application_tail_for_preparation(tmp_path: Path) -> None:
    folder = tmp_path / "raw" / "submissions" / "1 Alpha"
    folder.mkdir(parents=True)
    path = folder / "combined.docx"
    _write_article(path, "Alpha Author", "COMBINED ARTICLE TITLE")
    document = Document(path)
    document.add_paragraph("АНКЕТА УЧАСНИКА МІЖНАРОДНОЇ КОНФЕРЕНЦІЇ")
    document.add_paragraph("Ім’я і прізвище автора публікації")
    document.add_paragraph("Назва статті(або вільний слухач)")
    document.add_paragraph("Дата оплати організаційного внеску")
    document.save(path)

    manifest = build_article_manifest(tmp_path / "raw", tmp_path / "reports")

    assert manifest["manifest_status"] == "PASS"
    assert manifest["articles"][0]["provenance"]["embedded_service_tail"] is True
    assert "embedded_application_tail_requires_preparation" in manifest["articles"][0]["warnings"]


def _write_article(
    path: Path,
    author: str,
    title: str,
    *,
    udc: str = "УДК 001.1",
    body_token: str = "body",
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(udc)
    document.add_paragraph(author)
    document.add_paragraph(title)
    document.add_paragraph("Анотація: " + " ".join([body_token] * 180))
    document.add_paragraph("Ключові слова: automation, journal, evidence")
    document.add_paragraph(" ".join([body_token] * 220))
    document.add_paragraph("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ:")
    document.add_paragraph("1. Synthetic regression source.")
    document.save(path)


def _write_application(path: Path, author: str, title: str, section: str) -> None:
    document = Document()
    document.add_paragraph("АНКЕТА УЧАСНИКА")
    table = document.add_table(rows=2, cols=7)
    headers = [
        "№",
        "Ім’я і прізвище автора публікації",
        "Вч. звання",
        "Місце роботи",
        "Контактні дані",
        "Назва статті",
        "Секція",
    ]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
    values = ["1.", author, "", "", "", title, section]
    for index, value in enumerate(values):
        table.cell(1, index).text = value
    document.save(path)
