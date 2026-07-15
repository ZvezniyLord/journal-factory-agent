from __future__ import annotations

from pathlib import Path

from docx import Document

from journal_factory.production_preview import (
    _audit_frontmatter_metadata,
    _export_toc,
    _extract_identifiers,
    _identifier_audit,
    _insert_canonical_toc,
    _replace_frontmatter_text,
    _frontmatter_replacements,
    _ensure_preview_styles,
    ArticleCandidate,
)


def test_frontmatter_metadata_replacement_removes_previous_conference(tmp_path: Path) -> None:
    path = tmp_path / "frontmatter.docx"
    doc = Document()
    doc.add_paragraph("INFORMATION PLATFORM \"CENTER FOR INNOVATIVE THINKING\"")
    doc.add_paragraph("January 19-21, 2026")
    doc.add_paragraph("Cambridge, United Kingdom")
    doc.add_paragraph("https://naukainfo.com/conference?id=91")
    doc.add_paragraph("https://doi.org/10.64828/conf-91-2026")
    _replace_frontmatter_text(doc, _frontmatter_replacements())
    doc.save(path)

    audit = _audit_frontmatter_metadata(path, {})

    assert audit["institutional_header_present"] is False
    assert "Cambridge, United Kingdom" not in "\n".join(p.text for p in Document(path).paragraphs)
    assert "conf-91-2026" not in "\n".join(p.text for p in Document(path).paragraphs)


def test_doi_udc_orcid_identifier_audit_preserves_required_ids() -> None:
    candidate = _candidate(
        text="DOI: https://doi.org/10.64828/conf-138-2026-10\n"
        "УДК 792.02:001.4\n"
        "ORCID: 0009-0005-7848-4943\n"
        "ORCID: 0009-0009-4259-6467\n"
    )
    final = {
        "normalized_text": "DOI: https://doi.org/10.64828/conf-138-2026-10 УДК 792.02:001.4 "
        "ORCID: 0009-0005-7848-4943 ORCID: 0009-0009-4259-6467"
    }

    audit = _identifier_audit(candidate, final)

    assert audit["status"] == "REVIEW"
    assert audit["missing_doi"] == []
    assert audit["missing_orcids"] == []
    assert audit["doi_before_udc"] is True


def test_orcid_not_lost_when_email_removed() -> None:
    candidate = _candidate("Name email@example.com ORCID: 0000-0001-8157-7028")
    final = {"normalized_text": "Name ORCID: 0000-0001-8157-7028"}

    audit = _identifier_audit(candidate, final)

    assert audit["missing_orcids"] == []
    assert "REQUIRED_IDENTIFIER_LOST" not in audit["blockers"]


def test_missing_orcid_fails_with_required_identifier_lost() -> None:
    candidate = _candidate("ORCID: 0000-0001-8157-7028")
    final = {"normalized_text": "Author without identifier"}

    audit = _identifier_audit(candidate, final)

    assert audit["status"] == "FAIL"
    assert audit["missing_orcids"] == ["0000-0001-8157-7028"]
    assert "REQUIRED_IDENTIFIER_LOST" in audit["blockers"]


def test_toc_three_columns_export_has_no_decorative_rows(tmp_path: Path) -> None:
    path = tmp_path / "toc.docx"
    doc = Document()
    _ensure_preview_styles(doc)
    doc.add_paragraph("TABLE OF CONTENTS")
    _insert_canonical_toc(doc, [_candidate("DOI: 10.64828/conf-138-2026-1")])
    doc.save(path)

    audit = _export_toc(path, tmp_path / "toc.tsv")

    assert audit["status"] == "REVIEW"
    assert audit["column_count"] == 3
    assert audit["decorative_rows"] == []
    assert "\t" in (tmp_path / "toc.tsv").read_text(encoding="utf-8")


def test_identifier_counts_match_source_to_final() -> None:
    text = "DOI: 10.64828/conf-138-2026-1 УДК 001 ORCID: 0009-0009-9367-0604"
    candidate = _candidate(text)
    audit = _identifier_audit(candidate, {"normalized_text": text})

    assert len(audit["source_doi"]) == len(audit["final_doi"])
    assert len(audit["source_orcids"]) == len(audit["final_orcids"])
    assert len(audit["source_udc"]) == len(audit["final_udc"])


def _candidate(text: str) -> ArticleCandidate:
    return ArticleCandidate(
        article_id="article-test",
        source_path="fixture.docx",
        extracted_path=Path("fixture.docx"),
        source_sha256="abc",
        title="Fixture Title",
        authors="Fixture Author",
        has_udc=bool(_extract_identifiers(text)["udc"]),
        identifiers=_extract_identifiers(text),
        table_count=0,
        table_cells=[],
        media_hashes=[],
        textbox_text=[],
        paragraph_count=1,
        visible_units=[text],
        normalized_text=text,
        score=0,
    )
