from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

from journal_factory.article_preparation import prepare_article_source
from journal_factory.builder_fidelity import normalize_visible_text
from journal_factory.source_snapshot import extract_docx_evidence_text


def test_embedded_application_tail_is_removed_without_changing_article_prefix(tmp_path: Path) -> None:
    source = tmp_path / "combined.docx"
    document = Document()
    document.add_paragraph("AUTHOR")
    document.add_paragraph("ARTICLE TITLE")
    document.add_paragraph("IMMUTABLE ARTICLE BODY")
    document.add_paragraph("АНКЕТА УЧАСНИКА МІЖНАРОДНОЇ КОНФЕРЕНЦІЇ")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "PRIVATE APPLICATION DATA"
    document.save(source)

    prepared, report = prepare_article_source(
        source,
        {
            "article_id": "article-001",
            "source_path": "submissions/1 Author/combined.docx",
            "provenance": {"embedded_service_tail": True},
        },
        tmp_path / "prepared",
    )

    text = normalize_visible_text(extract_docx_evidence_text(prepared))
    assert report["status"] == "PASS"
    assert report["transformation"] == "remove_embedded_application_tail"
    assert report["preserved_article_text"] is True
    assert text == "AUTHOR ARTICLE TITLE IMMUTABLE ARTICLE BODY"
    assert "PRIVATE APPLICATION DATA" not in text


def test_article_without_service_tail_is_not_rewritten(tmp_path: Path) -> None:
    source = tmp_path / "article.docx"
    document = Document()
    document.add_paragraph("ARTICLE BODY")
    document.save(source)

    prepared, report = prepare_article_source(
        source,
        {
            "article_id": "article-001",
            "source_path": "article.docx",
            "provenance": {"embedded_service_tail": False},
        },
        tmp_path / "prepared",
    )

    assert prepared == source
    assert report["transformation"] == "none"
    assert report["source_sha256"] == report["prepared_sha256"]


def test_article_title_is_bookmarked_and_local_section_restart_is_removed(tmp_path: Path) -> None:
    source = tmp_path / "article-with-section.docx"
    document = Document()
    document.add_paragraph("ARTICLE TITLE")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("BODY")
    document.save(source)

    prepared, report = prepare_article_source(
        source,
        {
            "article_id": "article-007",
            "journal_order": 7,
            "source_path": source.name,
            "title_detected": "ARTICLE TITLE",
            "provenance": {"embedded_service_tail": False},
        },
        tmp_path / "prepared",
    )

    output = Document(prepared)
    assert report["status"] == "PASS"
    assert report["removed_section_properties"] == 1
    assert report["title_bookmark"] == "JF_ARTICLE_007_START"
    assert len(output.sections) == 1
    assert output.element.body.xpath(".//w:bookmarkStart[@w:name='JF_ARTICLE_007_START']")


def test_split_manifest_title_is_merged_without_changing_visible_text(tmp_path: Path) -> None:
    source = tmp_path / "split-title.docx"
    document = Document()
    document.add_paragraph("FIRST HALF OF")
    document.add_paragraph("THE ARTICLE TITLE")
    document.add_paragraph("BODY")
    document.save(source)
    original_text = normalize_visible_text(extract_docx_evidence_text(source))

    prepared, report = prepare_article_source(
        source,
        {
            "article_id": "article-008",
            "journal_order": 8,
            "source_path": source.name,
            "title_manifest": "FIRST HALF OF THE ARTICLE TITLE",
            "title_detected": "THE ARTICLE TITLE",
            "provenance": {"embedded_service_tail": False},
        },
        tmp_path / "prepared",
    )

    assert report["status"] == "PASS"
    assert report["title_match_source"] == "manifest"
    assert "merge_split_article_title" in report["transformation"]
    assert normalize_visible_text(extract_docx_evidence_text(prepared)) == original_text


def test_trailing_empty_paragraph_is_preserved_for_source_layout(tmp_path: Path) -> None:
    source = tmp_path / "trailing-empty.docx"
    document = Document()
    document.add_paragraph("ARTICLE TITLE")
    document.add_paragraph("BODY")
    document.add_paragraph("")
    document.save(source)

    prepared, report = prepare_article_source(
        source,
        {
            "article_id": "article-009",
            "journal_order": 9,
            "source_path": source.name,
            "title_detected": "ARTICLE TITLE",
            "provenance": {"embedded_service_tail": False},
        },
        tmp_path / "prepared",
    )

    assert report["status"] == "PASS"
    assert report["removed_trailing_empty_paragraphs"] == 0
    assert Document(prepared).paragraphs[-1].text == ""
