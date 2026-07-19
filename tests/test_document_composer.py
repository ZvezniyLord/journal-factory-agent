from __future__ import annotations

import base64
import io
from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from journal_factory.article_preparation import prepare_article_source
from journal_factory.document_composer import (
    _apply_layout_adjustments,
    _separator_document,
    _set_word_compatibility_mode,
    compose_articles_into_etalon,
    find_article_insert_index,
)


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZlFsAAAAASUVORK5CYII="
)


def test_article_is_inserted_after_toc_and_before_tail_with_objects(tmp_path: Path) -> None:
    master_path = tmp_path / "etalon.docx"
    article_path = tmp_path / "article.docx"
    output_path = tmp_path / "journal.docx"

    master = Document()
    master.styles.add_style("SECTION", WD_STYLE_TYPE.PARAGRAPH)
    master.add_paragraph("FRONT MATTER")
    master.add_paragraph("TABLE OF CONTENTS")
    master.add_section(WD_SECTION.NEW_PAGE)
    master.add_paragraph("FINAL SERVICE PAGE")
    master.save(master_path)

    assert find_article_insert_index(Document(master_path)) > 0

    article = Document()
    article.add_paragraph("UNIQUE ARTICLE BODY")
    article.add_table(rows=1, cols=1).cell(0, 0).text = "TABLE CELL"
    article.add_picture(io.BytesIO(ONE_PIXEL_PNG))
    article.save(article_path)

    report = compose_articles_into_etalon(
        master_path,
        [
            {
                "article_id": "article-001",
                "source_file": str(article_path),
                "section": "SECTION ONE",
            }
        ],
        output_path,
    )

    assert report["status"] == "PASS"
    assert report["inserted_count"] == 1
    output = Document(output_path)
    texts = [paragraph.text for paragraph in output.paragraphs]
    assert texts.index("TABLE OF CONTENTS") < texts.index("UNIQUE ARTICLE BODY")
    assert texts.index("UNIQUE ARTICLE BODY") < texts.index("FINAL SERVICE PAGE")
    first_boundary = next(
        item
        for item in output.element.body.xpath(".//w:bookmarkStart")
        if item.get(qn("w:name")) == "JF_LAYOUT_ARTICLE_001_START"
    ).getparent()
    assert first_boundary.xpath("./w:pPr/w:pageBreakBefore")
    assert len(output.tables) == 1
    assert output.tables[0].cell(0, 0).text == "TABLE CELL"

    with zipfile.ZipFile(output_path) as package:
        media = [name for name in package.namelist() if name.startswith("word/media/")]
    assert media


def test_missing_toc_anchor_fails_closed(tmp_path: Path) -> None:
    master_path = tmp_path / "bad-etalon.docx"
    master = Document()
    master.add_paragraph("NO TOC")
    master.save(master_path)

    try:
        find_article_insert_index(Document(master_path))
    except ValueError as exc:
        assert "TABLE OF CONTENTS" in str(exc)
    else:
        raise AssertionError("Missing TOC anchor must fail closed")


def test_separator_uses_page_break_before_without_creating_a_blank_page() -> None:
    separator = _separator_document(True, "CANONICAL SECTION")
    paragraph = separator.paragraphs[0]

    assert paragraph.text == "CANONICAL SECTION"
    assert paragraph._p.xpath("./w:pPr/w:pageBreakBefore")
    assert not paragraph._p.xpath(".//w:br[@w:type='page']")


def test_word_compatibility_mode_is_explicit_and_replaces_existing_value() -> None:
    document = Document()
    _set_word_compatibility_mode(document, 14)
    _set_word_compatibility_mode(document, 14)

    settings = document.settings.element.xpath(
        ".//w:compatSetting[@w:name='compatibilityMode']"
    )
    assert len(settings) == 1
    assert settings[0].get(qn("w:val")) == "14"


def test_layout_adjustments_apply_content_anchored_editorial_geometry() -> None:
    document = Document()
    document.styles.add_style("Normal (Web)", WD_STYLE_TYPE.PARAGRAPH)
    boundary = document.add_paragraph()
    boundary._p.get_or_add_pPr().append(OxmlElement("w:pageBreakBefore"))
    boundary_start = OxmlElement("w:bookmarkStart")
    boundary_start.set(qn("w:id"), "4002")
    boundary_start.set(qn("w:name"), "JF_LAYOUT_ARTICLE_002_START")
    boundary._p.append(boundary_start)
    boundary_end = OxmlElement("w:bookmarkEnd")
    boundary_end.set(qn("w:id"), "4002")
    boundary._p.append(boundary_end)
    document.add_paragraph("UDC 001")
    document.add_paragraph("")
    document.add_paragraph("")
    author = document.add_paragraph("        AUTHOR NAME")
    author.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("")
    document.add_paragraph("")
    title = document.add_paragraph("ARTICLE TITLE")
    bookmark = OxmlElement("w:bookmarkStart")
    bookmark.set(qn("w:id"), "2")
    bookmark.set(qn("w:name"), "JF_ARTICLE_002_START")
    title._p.insert(0, bookmark)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "2")
    title._p.append(bookmark_end)
    body = document.add_paragraph("BODY TEXT")
    body.style = "Normal (Web)"
    body.paragraph_format.space_before = Pt(12)
    body.paragraph_format.space_after = Pt(12)
    character_spacing = OxmlElement("w:spacing")
    character_spacing.set(qn("w:val"), "15")
    body.runs[0]._r.get_or_add_rPr().append(character_spacing)
    anchor = document.add_paragraph("CONTENT ANCHOR CONTINUES")
    following = document.add_paragraph()
    following._p.get_or_add_pPr().append(OxmlElement("w:pageBreakBefore"))
    following_start = OxmlElement("w:bookmarkStart")
    following_start.set(qn("w:id"), "4003")
    following_start.set(qn("w:name"), "JF_LAYOUT_ARTICLE_003_START")
    following._p.append(following_start)
    following_end = OxmlElement("w:bookmarkEnd")
    following_end.set(qn("w:id"), "4003")
    following._p.append(following_end)

    report = _apply_layout_adjustments(
        document,
        [
            {
                "ordinal": 2,
                "trim_right_aligned_padding": True,
                "collapse_redundant_pre_title_empty_paragraphs": True,
                "paragraph_styles_to_normal": ["Normal (Web)"],
                "normalize_paragraph_spacing": True,
                "normalize_character_spacing": True,
                "page_break_before_anchor": "CONTENT ANCHOR",
            }
        ],
    )

    assert report["status"] == "PASS"
    applied = report["applied"][0]
    assert applied["trimmed_right_aligned_padding_paragraphs"] == 1
    assert applied["collapsed_pre_title_empty_paragraphs"] == 2
    assert applied["paragraph_styles_normalized"] == 1
    assert applied["paragraph_spacings_normalized"] == 1
    assert applied["page_break_inserted"] is True
    assert author.text == "AUTHOR NAME"
    assert body.style.name == "Normal"
    assert body._p.find("./" + qn("w:pPr") + "/" + qn("w:spacing")).get(
        qn("w:before")
    ) == "0"
    assert body.runs[0]._r.find("./" + qn("w:rPr") + "/" + qn("w:spacing")).get(
        qn("w:val")
    ) == "0"
    assert anchor._p.xpath(".//w:br[@w:type='page']")


def test_composer_creates_custom_style_toc_and_continuous_page_numbers(tmp_path: Path) -> None:
    master_path = tmp_path / "etalon.docx"
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "journal.docx"
    master = Document()
    master.styles.add_style("SECTION", WD_STYLE_TYPE.PARAGRAPH)
    master.styles.add_style("Назва1", WD_STYLE_TYPE.PARAGRAPH)
    master.add_paragraph("FRONT MATTER")
    master.add_paragraph("TABLE OF CONTENTS")
    master.add_section(WD_SECTION.NEW_PAGE)
    master.add_paragraph("FINAL SERVICE PAGE")
    master.save(master_path)
    source = Document()
    source.add_paragraph("ARTICLE TITLE")
    source.add_paragraph("BODY")
    source.save(source_path)
    prepared, preparation = prepare_article_source(
        source_path,
        {
            "article_id": "article-001",
            "journal_order": 1,
            "source_path": source_path.name,
            "title_detected": "ARTICLE TITLE",
            "provenance": {"embedded_service_tail": False},
        },
        tmp_path / "prepared",
    )

    report = compose_articles_into_etalon(
        master_path,
        [
            {
                "article_id": "article-001",
                "source_file": str(prepared),
                "section": "PEDAGOGY AND EDUCATION",
                "title_bookmark": preparation["title_bookmark"],
            }
        ],
        output_path,
    )

    assert report["status"] == "PASS"
    assert report["structure"]["title_styles_applied"] == 1
    output = Document(output_path)
    title = next(paragraph for paragraph in output.paragraphs if paragraph.text == "ARTICLE TITLE")
    assert title.style.name == "Назва1"
    instructions = " ".join(output.element.body.xpath(".//w:instrText/text()"))
    assert "TOC \\h \\z \\t" in instructions
    assert "SECTION" in instructions
    assert "Назва1" in instructions
    assert report["structure"]["toc_style_mapping"] == {"SECTION": 1, "Назва1": 2}
    section_properties = output.element.body.xpath(".//w:sectPr")
    assert section_properties[0].find(qn("w:pgNumType")).get(qn("w:start")) == "1"
    assert all(item.find(qn("w:pgNumType")) is None for item in section_properties[1:])
    assert output.element.body.xpath(".//w:footerReference")
