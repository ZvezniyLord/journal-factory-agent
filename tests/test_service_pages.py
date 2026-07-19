from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION

from journal_factory.conference_metadata import load_conference_metadata
from journal_factory.service_pages import materialize_service_pages


def test_front_matter_replacement_handles_split_runs_and_dynamic_page_field(tmp_path: Path) -> None:
    source = tmp_path / "etalon.docx"
    output = tmp_path / "materialized.docx"
    config = tmp_path / "conference.json"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("OLD ")
    paragraph.add_run("EVENT")
    document.add_paragraph("OLD CITY")
    document.add_paragraph("https://example.test/conference?id=91")
    document.add_paragraph("0000000 p.")
    document.save(source)
    config.write_text(
        """{
          "schema_version": 1,
          "conference_id": 95,
          "event_title": "New Event",
          "dates": "February 6-8, 2026",
          "location": "Oxford, United Kingdom",
          "conference_url": "https://example.test/conference?id=95",
          "doi_url": "https://doi.org/10.1/conf-95-2026",
          "isbn": "978-0-00-000095-0",
          "udc": "001",
          "typography_profile": "legacy_14pt",
          "template_replacements": [
            {"source": "OLD EVENT", "target": "{event_title}"},
            {"source": "OLD CITY", "target": "{location}"},
            {"source": "conference?id=91", "target": "conference?id=95"}
          ],
          "dynamic_fields": [{"placeholder": "0000000", "field": "NUMPAGES"}],
          "required_markers": ["{event_title}", "{location}", "conference?id=95"],
          "stale_markers": ["OLD EVENT", "OLD CITY", "conference?id=91"],
          "provenance": [{"field": "all", "source": "fixture", "confidence": 1.0}]
        }""",
        encoding="utf-8",
    )

    report = materialize_service_pages(source, output, load_conference_metadata(config))

    assert report["status"] == "PASS"
    assert report["audit"]["stale_marker_count"] == 0
    assert Document(output).paragraphs[0].text == "New Event"
    with zipfile.ZipFile(output) as package:
        xml = package.read("word/document.xml").decode("utf-8")
    assert 'w:instr="NUMPAGES"' in xml


def test_conference_95_metadata_matches_schema_contract() -> None:
    metadata = load_conference_metadata(Path("fixtures/conferences/095/conference_metadata.json"))
    assert metadata["conference_id"] == 95
    assert metadata["event_title"] == "Oxford International Science Forum"
    assert metadata["typography_profile"] == "legacy_14pt"


def test_service_materialization_trims_only_unprotected_trailing_empty_paragraphs(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("OLD EVENT")
    document.add_paragraph("OLD CITY")
    document.add_paragraph("")
    document.save(source)
    metadata = {
        "conference_id": 95,
        "event_title": "New Event",
        "location": "Oxford",
        "template_replacements": [
            {"source": "OLD EVENT", "target": "{event_title}"},
            {"source": "OLD CITY", "target": "{location}"},
        ],
        "required_markers": ["{event_title}", "{location}"],
        "stale_markers": ["OLD EVENT", "OLD CITY"],
        "service_layout": {"trim_trailing_empty_paragraphs": True},
    }

    report = materialize_service_pages(source, output, metadata)

    assert report["status"] == "PASS"
    assert report["layout_adjustment_counts"]["trailing_empty_paragraphs"] >= 1
    assert Document(output).paragraphs[-1].text == "Oxford"


def test_service_materialization_builds_locked_official_toc_and_special_thanks(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("OLD EVENT")
    document.add_paragraph("TABLE OF CONTENTS")
    document.add_paragraph("placeholder")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("tail")
    document.save(source)
    metadata = {
        "conference_id": 95,
        "event_title": "New Event",
        "template_replacements": [{"source": "OLD EVENT", "target": "{event_title}"}],
        "required_markers": ["{event_title}"],
        "stale_markers": ["OLD EVENT"],
        "official_toc": {
            "toc_layout": {
                "line_spacing_twips": 240,
                "special_thanks_line_spacing_twips": 240,
            },
            "articles": [
                {
                    "ordinal": 1,
                    "section": "ECONOMIC THEORY",
                    "authors_display": "Roman Reznikov",
                    "title": "FIRST ARTICLE TITLE",
                    "printed_start_page": 6,
                }
            ],
            "special_thanks": {
                "present": True,
                "heading": "SPECIAL THANKS TO PARTICIPANTS:",
                "participants_display": "One Person, Two Person",
            },
        },
    }

    report = materialize_service_pages(source, output, metadata)

    assert report["status"] == "PASS"
    assert report["layout_adjustment_counts"]["official_toc_articles"] == 1
    assert report["layout_adjustment_counts"]["official_toc_sections"] == 1
    assert report["audit"]["special_thanks_present"] is True
    with zipfile.ZipFile(output) as package:
        xml = package.read("word/document.xml").decode("utf-8")
    assert 'w:fldLock="true"' in xml
    assert "PAGEREF JF_ARTICLE_001_START" in xml
    assert "SPECIAL THANKS TO PARTICIPANTS:" in xml
    assert 'w:name="JF_SPECIAL_THANKS"' in xml
    assert 'w:line="240"' in xml
