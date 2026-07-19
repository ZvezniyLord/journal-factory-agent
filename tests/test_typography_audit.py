from pathlib import Path
import json

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from journal_factory.article_preparation import prepare_article_source
from journal_factory.typography_audit import audit_effective_typography
from journal_factory.typography_profile import apply_typography_profile


def test_effective_typography_audits_legacy_and_standard_profiles(tmp_path: Path) -> None:
    profiles = tmp_path / "profiles.json"
    profiles.write_text(
        json.dumps(
            {
                "profiles": {
                    "legacy_14pt": {"body_font_size_pt": 14},
                    "standard_11pt": {"body_font_size_pt": 11},
                }
            }
        ),
        encoding="utf-8",
    )
    for profile, expected in (("legacy_14pt", 14), ("standard_11pt", 11)):
        source = tmp_path / f"{profile}-source.docx"
        prepared_dir = tmp_path / f"{profile}-prepared"
        profiled = tmp_path / f"{profile}-profiled.docx"
        document = Document()
        document.add_paragraph("ARTICLE TITLE")
        document.add_paragraph("Body inheriting the configured default size")
        paragraph_mark = document.add_paragraph("Paragraph mark size is not a text-run size")
        paragraph_rpr = OxmlElement("w:rPr")
        paragraph_size = OxmlElement("w:sz")
        paragraph_size.set(qn("w:val"), str((expected + 4) * 2))
        paragraph_rpr.append(paragraph_size)
        paragraph_mark._p.get_or_add_pPr().append(paragraph_rpr)
        override = document.add_paragraph()
        override.add_run("Explicit author override").font.size = Pt(expected - 2)
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).paragraphs[0].add_run("Compact table text").font.size = Pt(8)
        document.save(source)
        prepared, preparation = prepare_article_source(
            source,
            {
                "article_id": "article-001",
                "journal_order": 1,
                "source_path": source.name,
                "title_detected": "ARTICLE TITLE",
                "provenance": {"embedded_service_tail": False},
            },
            prepared_dir,
        )
        assert preparation["status"] == "PASS"
        apply_typography_profile(prepared, profiled, profiles, profile)

        report = audit_effective_typography(profiled, profiles, profile)

        assert report["status"] == "PASS"
        assert report["expected_body_font_size_pt"] == expected
        assert report["effective_body_size_counts"] == {str(expected): 3}
        assert report["paragraph_mark_size_counts"] == {str(expected + 4): 1}
        assert report["exception_counts"]["local_author_override"] == 1
